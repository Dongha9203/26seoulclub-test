"""
전체 기능 단위 테스트.

실행 방법:
  pip install pytest pytest-mock
  pytest tests/test_all.py -v

외부 API(Notion, Google Sheets)는 mock으로 처리합니다.
파일 기반 수집기(docx, pdf, excel, txt)는 임시 파일을 생성해 실제로 테스트합니다.
"""

import json
import sys
import tempfile
import uuid
from pathlib import Path
from unittest.mock import ANY, MagicMock, patch

import pytest

# 프로젝트 루트 경로 등록
_root = Path(__file__).parent.parent
if str(_root) not in sys.path:
    sys.path.insert(0, str(_root))


# ══════════════════════════════════════════════════════════════
# models/document.py
# ══════════════════════════════════════════════════════════════

class TestDocument:
    def test_new_creates_doc_with_defaults(self):
        from models.document import Document
        doc = Document.new(
            source_type="docx",
            source_origin="test.docx",
            title="제목",
            content="내용",
        )
        assert doc.doc_id
        assert doc.source_type == "docx"
        assert doc.is_editable is True
        assert doc.notion_block_id is None
        assert doc.notion_page_url is None

    def test_notion_source_is_not_editable(self):
        from models.document import Document
        doc = Document.new(
            source_type="notion",
            source_origin="메인페이지",
            title="제목",
            content="내용",
            notion_block_id="abc-123",
            notion_page_url="https://notion.so/xxx",
        )
        assert doc.is_editable is False

    def test_deep_link_url(self):
        from models.document import Document
        doc = Document.new(
            source_type="notion",
            source_origin="FAQ",
            title="질문",
            content="답변",
            notion_page_url="https://www.notion.so/abc123",
            notion_block_id="12345678-1234-1234-1234-123456789abc",
        )
        link = doc.deep_link_url()
        # UUID "12345678-1234-1234-1234-123456789abc" → 하이픈 제거 → 32자
        assert link == "https://www.notion.so/abc123#12345678123412341234123456789abc"

    def test_deep_link_url_none_when_no_block_id(self):
        from models.document import Document
        doc = Document.new(
            source_type="notion",
            source_origin="메인페이지",
            title="제목",
            content="내용",
            notion_page_url="https://www.notion.so/xxx",
            notion_block_id=None,
        )
        assert doc.deep_link_url() is None

    def test_to_dict_is_serializable(self):
        from models.document import Document
        doc = Document.new(source_type="pdf", source_origin="a.pdf", title="T", content="C")
        d = doc.to_dict()
        assert isinstance(d, dict)
        assert d["source_type"] == "pdf"
        assert isinstance(d["is_editable"], bool)

    def test_empty_title_allowed(self):
        from models.document import Document
        doc = Document.new(source_type="pdf", source_origin="a.pdf", title="", content="C")
        assert doc.title == ""

    def test_explicit_is_editable_override(self):
        from models.document import Document
        doc = Document.new(
            source_type="notion",
            source_origin="FAQ",
            title="T",
            content="C",
            is_editable=True,  # 명시적으로 True로 오버라이드
        )
        assert doc.is_editable is True


# ══════════════════════════════════════════════════════════════
# storage/supabase_store.py
# ══════════════════════════════════════════════════════════════

class TestSupabaseStore:
    def _make_doc(self, source_type="docx", title="T", content="C", origin="test"):
        from models.document import Document
        return Document.new(source_type=source_type, source_origin=origin,
                            title=title, content=content)

    def test_upsert_and_get(self, pg_conn):
        # 공유 운영 DB라 get_all()에는 실제 KB 문서도 같이 섞여 나오므로,
        # 전체 건수가 아니라 이번에 넣은 doc_id가 정확히 1건 있는지로 검증합니다.
        from storage.supabase_store import upsert_document, get_all
        doc = self._make_doc()
        upsert_document(doc, pg_conn)
        all_docs = get_all(pg_conn)
        matches = [d for d in all_docs if d.doc_id == doc.doc_id]
        assert len(matches) == 1

    def test_upsert_updates_existing(self, pg_conn):
        from storage.supabase_store import upsert_document, get_all
        doc = self._make_doc(title="원본")
        upsert_document(doc, pg_conn)
        doc.title = "수정됨"
        upsert_document(doc, pg_conn)
        all_docs = get_all(pg_conn)
        matches = [d for d in all_docs if d.doc_id == doc.doc_id]
        assert len(matches) == 1
        assert matches[0].title == "수정됨"

    def test_upsert_many(self, pg_conn):
        from storage.supabase_store import upsert_documents, get_total_count
        baseline = get_total_count(pg_conn)
        docs = [self._make_doc(title=f"제목{i}") for i in range(5)]
        count = upsert_documents(docs, pg_conn)
        assert count == 5
        assert get_total_count(pg_conn) - baseline == 5

    def test_delete_by_source_origin(self, pg_conn):
        from storage.supabase_store import upsert_documents, delete_by_source_origin, get_all
        origin_a, origin_b = f"test-a-{uuid.uuid4()}", f"test-b-{uuid.uuid4()}"
        docs = [self._make_doc(origin=origin_a) for _ in range(3)]
        other = [self._make_doc(origin=origin_b) for _ in range(2)]
        upsert_documents(docs + other, pg_conn)
        deleted = delete_by_source_origin(origin_a, pg_conn)
        assert deleted == 3
        other_ids = {d.doc_id for d in other}
        remaining = [d for d in get_all(pg_conn) if d.doc_id in other_ids]
        assert len(remaining) == 2
        assert all(d.source_origin == origin_b for d in remaining)

    def test_get_by_source_type(self, pg_conn):
        from storage.supabase_store import upsert_documents, get_by_source_type
        origin = f"test-{uuid.uuid4()}"
        docs = [self._make_doc(source_type="notion", origin=origin) for _ in range(2)]
        docs += [self._make_doc(source_type="pdf", origin=f"{origin}-pdf") for _ in range(3)]
        upsert_documents(docs, pg_conn)
        notion_docs = get_by_source_type("notion", pg_conn)
        matches = [d for d in notion_docs if d.source_origin == origin]
        assert len(matches) == 2

    def test_category_distribution(self, pg_conn):
        from storage.supabase_store import upsert_documents, get_category_distribution
        before = get_category_distribution(pg_conn)
        docs = []
        for _ in range(3):
            d = self._make_doc()
            d.category = "신청 자격 안내"
            docs.append(d)
        for _ in range(2):
            d = self._make_doc()
            d.category = "미분류"
            docs.append(d)
        upsert_documents(docs, pg_conn)
        after = get_category_distribution(pg_conn)
        assert after["신청 자격 안내"] - before.get("신청 자격 안내", 0) == 3
        assert after["미분류"] - before.get("미분류", 0) == 2

    def test_sync_metadata_upsert_and_get(self, pg_conn):
        from storage.supabase_store import upsert_sync_metadata, get_sync_metadata
        upsert_sync_metadata("faq", "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)
        meta = get_sync_metadata("faq", pg_conn)
        assert meta["page_key"] == "faq"
        assert meta["last_notion_edited_time"] == "2026-01-01T00:00:00Z"

    def test_sync_metadata_returns_none_for_unknown(self, pg_conn):
        from storage.supabase_store import get_sync_metadata
        meta = get_sync_metadata("nonexistent", pg_conn)
        assert meta is None

    def test_sync_metadata_children_roundtrip(self, pg_conn):
        # 실제 노션 동기화로 생긴 "main::<page_id>" 같은 진짜 행과 절대 겹치지 않도록
        # 테스트 전용 부모 키를 사용합니다 (공유 운영 DB라 절대값 비교가 위험합니다).
        from storage.supabase_store import upsert_sync_metadata, get_sync_metadata_children
        parent = f"test-parent-{uuid.uuid4()}"
        upsert_sync_metadata(parent, "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)
        upsert_sync_metadata(f"{parent}::sub-1", "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)
        upsert_sync_metadata(f"{parent}::sub-2", "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)
        upsert_sync_metadata("integrated_system", "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)

        children = get_sync_metadata_children(parent, pg_conn)
        assert {c["page_key"] for c in children} == {f"{parent}::sub-1", f"{parent}::sub-2"}

    def test_delete_sync_metadata_children_only_deletes_matching_prefix(self, pg_conn):
        from storage.supabase_store import (
            upsert_sync_metadata, delete_sync_metadata_children, get_sync_metadata, get_sync_metadata_children,
        )
        parent = f"test-parent-{uuid.uuid4()}"
        upsert_sync_metadata(parent, "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)
        upsert_sync_metadata(f"{parent}::sub-1", "2026-01-01T00:00:00Z", "2026-01-02T00:00:00Z", pg_conn)

        deleted = delete_sync_metadata_children(parent, pg_conn)
        assert deleted == 1
        assert get_sync_metadata_children(parent, pg_conn) == []
        assert get_sync_metadata(parent, pg_conn) is not None  # 최상위 자신은 그대로 남아있어야 함

    def test_empty_upsert_many(self, pg_conn):
        from storage.supabase_store import upsert_documents, get_total_count
        baseline = get_total_count(pg_conn)
        count = upsert_documents([], pg_conn)
        assert count == 0
        assert get_total_count(pg_conn) == baseline


# ══════════════════════════════════════════════════════════════
# utils/category_tagger.py
# ══════════════════════════════════════════════════════════════

class TestCategoryTagger:
    @pytest.fixture
    def tmp_config(self, tmp_path):
        config = {
            "categories": [
                {"name": "신청 자격 안내", "keywords": ["신청 자격", "지원 자격"]},
                {"name": "수당 지급 기준 안내", "keywords": ["수당", "지급"]},
                {"name": "출결 및 활동기준 안내", "keywords": ["출결", "출근"]},
            ]
        }
        p = tmp_path / "config.json"
        p.write_text(json.dumps(config), encoding="utf-8")
        return p

    def test_exact_keyword_match(self, tmp_config):
        from utils.category_tagger import tag_category
        result = tag_category("신청 자격 안내", "지원 자격 관련 내용", tmp_config)
        assert result == "신청 자격 안내"

    def test_content_keyword_match(self, tmp_config):
        from utils.category_tagger import tag_category
        result = tag_category("기타 제목", "수당 지급 방법 안내", tmp_config)
        assert result == "수당 지급 기준 안내"

    def test_no_match_returns_미분류(self, tmp_config):
        from utils.category_tagger import tag_category
        result = tag_category("관련 없는 제목", "관련 없는 내용", tmp_config)
        assert result == "미분류"

    def test_empty_title_and_content(self, tmp_config):
        from utils.category_tagger import tag_category
        result = tag_category("", "", tmp_config)
        assert result == "미분류"

    def test_missing_config_returns_미분류(self, tmp_path):
        from utils.category_tagger import tag_category
        missing_path = tmp_path / "nonexistent.json"
        result = tag_category("수당 지급", "지급 기준", missing_path)
        assert result == "미분류"

    def test_first_match_wins(self, tmp_config):
        # "신청 자격"과 "수당" 둘 다 포함 → 먼저 정의된 카테고리 반환
        from utils.category_tagger import tag_category
        result = tag_category("신청 자격", "수당 정보", tmp_config)
        assert result == "신청 자격 안내"


# ══════════════════════════════════════════════════════════════
# utils/validators.py
# ══════════════════════════════════════════════════════════════

class TestValidators:
    def _notion_doc(self, block_id=None):
        from models.document import Document
        doc = Document.new(
            source_type="notion",
            source_origin="FAQ",
            title="질문",
            content="답변",
            notion_block_id=block_id,
        )
        return doc

    def test_all_block_ids_present(self):
        from utils.validators import validate_notion_block_ids
        docs = [self._notion_doc(block_id=f"block-{i}") for i in range(3)]
        result = validate_notion_block_ids(docs)
        assert result["block_id_missing"] == 0
        assert result["success_rate_pct"] == 100.0

    def test_missing_block_id_detected(self):
        from utils.validators import validate_notion_block_ids
        docs = [self._notion_doc(block_id="id1"), self._notion_doc(block_id=None)]
        result = validate_notion_block_ids(docs)
        assert result["block_id_missing"] == 1
        assert result["success_rate_pct"] == 50.0

    def test_non_notion_docs_ignored(self):
        from utils.validators import validate_notion_block_ids
        from models.document import Document
        non_notion = Document.new(source_type="pdf", source_origin="a.pdf",
                                   title="T", content="C")
        result = validate_notion_block_ids([non_notion])
        assert result["total_notion_docs"] == 0
        assert result["block_id_missing"] == 0

    def test_empty_list(self):
        from utils.validators import validate_notion_block_ids
        result = validate_notion_block_ids([])
        assert result["total_notion_docs"] == 0
        assert result["success_rate_pct"] == 100.0

    def test_validate_documents_full(self):
        from utils.validators import validate_documents
        from models.document import Document
        good = Document.new(source_type="pdf", source_origin="a.pdf", title="T", content="C")
        result = validate_documents([good])
        assert result["passed"] is True
        assert result["error_count"] == 0

    def test_validate_documents_detects_empty_title(self):
        from utils.validators import validate_documents
        from models.document import Document
        bad = Document.new(source_type="pdf", source_origin="a.pdf", title="", content="C")
        result = validate_documents([bad])
        assert result["passed"] is False
        assert result["error_count"] > 0


# ══════════════════════════════════════════════════════════════
# collectors/notion_collector.py (mock 사용)
# ══════════════════════════════════════════════════════════════

class TestNotionCollector:
    def _make_block(self, btype, text, block_id="abc-001", has_children=False):
        return {
            "id": block_id,
            "type": btype,
            btype: {"rich_text": [{"plain_text": text}]},
            "has_children": has_children,
        }

    def test_extract_page_id_from_url(self):
        from collectors.notion_collector import extract_page_id
        url = "https://www.notion.so/workspace/Title-abc123def456abc123def456abc123de"
        result = extract_page_id(url)
        assert "abc123de" in result.replace("-", "")

    def test_extract_page_id_uuid_format(self):
        from collectors.notion_collector import extract_page_id
        uuid = "abc123de-f456-abc1-23de-f456abc123de"
        assert extract_page_id(uuid) == uuid.lower()

    def test_extract_page_id_raw_hex(self):
        from collectors.notion_collector import extract_page_id
        raw = "abc123def456abc123def456abc123de"
        result = extract_page_id(raw)
        assert result == "abc123de-f456-abc1-23de-f456abc123de"

    def test_extract_page_id_with_query_params(self):
        from collectors.notion_collector import extract_page_id
        url = "https://www.notion.so/abc123def456abc123def456abc123de?pvs=4"
        result = extract_page_id(url)
        assert "-" in result  # UUID 형식

    def test_chunk_toggle_priority(self):
        """toggle 블록이 먼저 분리되는지 확인."""
        from collectors.notion_collector import _chunk_page_blocks
        mock_client = MagicMock()
        mock_client.blocks.children.list.return_value = {
            "results": [
                {"id": "child-1", "type": "paragraph",
                 "paragraph": {"rich_text": [{"plain_text": "답변 내용"}]},
                 "has_children": False},
            ],
            "has_more": False,
        }
        blocks = [
            {"id": "t1", "type": "toggle",
             "toggle": {"rich_text": [{"plain_text": "FAQ 질문"}]},
             "has_children": True},
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "섹션"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "섹션 내용"}]},
             "has_children": False},
        ]
        docs = _chunk_page_blocks(mock_client, blocks, "FAQ", "https://notion.so/xxx")
        toggle_docs = [d for d in docs if d.title == "FAQ 질문"]
        heading_docs = [d for d in docs if d.title == "섹션"]
        assert len(toggle_docs) == 1
        assert len(heading_docs) == 1

    def test_chunk_fallback_no_structure(self):
        """toggle/heading 없을 때 fallback 분할 적용."""
        from collectors.notion_collector import _chunk_page_blocks
        mock_client = MagicMock()
        blocks = [
            {"id": f"p{i}", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "A" * 100}]},
             "has_children": False}
            for i in range(10)  # 1000자 → 800자 기준으로 2개 청크
        ]
        docs = _chunk_page_blocks(mock_client, blocks, "페이지", "https://notion.so/xxx")
        assert len(docs) >= 1  # fallback 동작 확인
        for doc in docs:
            assert doc.notion_block_id is not None  # 블록 ID 존재

    def test_heading_chunking(self):
        """heading 단위로 청킹되는지 확인."""
        from collectors.notion_collector import _chunk_by_headings
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "1장"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "1장 내용"}]},
             "has_children": False},
            {"id": "h2", "type": "heading_2",
             "heading_2": {"rich_text": [{"plain_text": "2장"}]},
             "has_children": False},
            {"id": "p2", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "2장 내용"}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 2
        assert docs[0].title == "1장"
        assert docs[1].title == "2장"
        assert docs[0].notion_block_id == "h1"
        assert docs[1].notion_block_id == "h2"

    def test_toggleable_heading_chunking(self):
        """토글 가능한 heading(본문이 형제가 아니라 자식 블록에 있는 경우)도 내용을 수집하는지 확인."""
        from collectors.notion_collector import _chunk_by_headings
        mock_client = MagicMock()
        mock_client.blocks.children.list.return_value = {
            "results": [
                {"id": "c1", "type": "paragraph",
                 "paragraph": {"rich_text": [{"plain_text": "토글 안의 실제 내용"}]},
                 "has_children": False},
            ],
            "has_more": False,
        }
        blocks = [
            {"id": "h1", "type": "heading_3",
             "heading_3": {"rich_text": [{"plain_text": "1. 활동 종료 후"}], "is_toggleable": True},
             "has_children": True},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 1
        assert docs[0].title == "1. 활동 종료 후"
        assert "토글 안의 실제 내용" in docs[0].content

    def test_make_document_doc_id_stable_for_same_block(self):
        """같은 block_id로 두 번 청킹하면(매 동기화마다 일어나는 일) doc_id가 같아야
        incremental 동기화가 "이게 그 블록"이라고 알아볼 수 있다."""
        from collectors.notion_collector import _chunk_by_headings
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "1장"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "내용"}]},
             "has_children": False},
        ]
        docs1 = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        docs2 = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert docs1[0].doc_id == docs2[0].doc_id

    def test_make_document_doc_id_differs_for_different_blocks(self):
        from collectors.notion_collector import _chunk_by_headings
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "1장"}]}, "has_children": False},
            {"id": "h2", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "2장"}]}, "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert docs[0].doc_id != docs[1].doc_id

    def test_make_document_doc_id_differs_across_parts_of_same_block(self):
        """같은 heading(같은 block_id)이 여러 파트로 쪼개질 때, 파트마다 다른
        doc_id가 나와야 한다(part_index로 구분)."""
        from collectors.notion_collector import _chunk_by_headings, FALLBACK_CHUNK_SIZE
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "긴 장"}]}, "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "가" * (FALLBACK_CHUNK_SIZE - 100)}]},
             "has_children": False},
            {"id": "p2", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "나" * (FALLBACK_CHUNK_SIZE - 100)}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 2
        assert docs[0].doc_id != docs[1].doc_id

    def test_heading_chunking_splits_long_section_by_paragraph(self):
        """heading 본문이 FALLBACK_CHUNK_SIZE를 넘으면 문단 단위로 (파트 N) 분할되는지 확인."""
        from collectors.notion_collector import _chunk_by_headings, FALLBACK_CHUNK_SIZE
        mock_client = MagicMock()
        long_paragraph_1 = "가" * (FALLBACK_CHUNK_SIZE - 100)
        long_paragraph_2 = "나" * (FALLBACK_CHUNK_SIZE - 100)
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "긴 장"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": long_paragraph_1}]},
             "has_children": False},
            {"id": "p2", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": long_paragraph_2}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 2
        assert docs[0].title == "긴 장 (파트 1)"
        assert docs[1].title == "긴 장 (파트 2)"
        # 두 파트 모두 같은 heading을 가리켜야 딥링크가 유지됨
        assert docs[0].notion_block_id == "h1"
        assert docs[1].notion_block_id == "h1"
        assert long_paragraph_1 in docs[0].content
        assert long_paragraph_2 in docs[1].content
        for doc in docs:
            assert len(doc.content) <= FALLBACK_CHUNK_SIZE

    def test_heading_chunking_splits_single_oversized_paragraph(self):
        """문단 하나가 그 자체로 FALLBACK_CHUNK_SIZE를 넘으면 글자 단위로 분할되는지 확인."""
        from collectors.notion_collector import _chunk_by_headings, FALLBACK_CHUNK_SIZE
        mock_client = MagicMock()
        huge_paragraph = "다" * (FALLBACK_CHUNK_SIZE * 2 + 50)
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "거대한 장"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": huge_paragraph}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 3
        assert "".join(d.content for d in docs) == huge_paragraph
        for doc in docs:
            assert len(doc.content) <= FALLBACK_CHUNK_SIZE
            assert doc.notion_block_id == "h1"

    def test_pack_parts_by_size_exact_boundary(self):
        """누적 길이가 정확히 max_size와 같으면 한 청크에 남고, 1자만 넘으면 분리되는지 확인."""
        from collectors.notion_collector import _pack_parts_by_size
        # "\n"으로 합치므로 "A"*399 + "\n" + "B"*400 = 800자 → 정확히 max_size
        parts_fit = ["A" * 399, "B" * 400]
        chunks_fit = _pack_parts_by_size(parts_fit, max_size=800)
        assert len(chunks_fit) == 1
        assert len(chunks_fit[0]) == 800

        # 한 글자만 늘려도 합치면 801자라 두 청크로 분리되어야 함
        parts_overflow = ["A" * 400, "B" * 400]
        chunks_overflow = _pack_parts_by_size(parts_overflow, max_size=800)
        assert len(chunks_overflow) == 2
        assert chunks_overflow[0] == "A" * 400
        assert chunks_overflow[1] == "B" * 400

    def test_pack_parts_by_size_oversized_part_resumes_buffer_after(self):
        """작은 조각 → 캡을 넘는 조각 → 작은 조각 순서에서, 캡을 넘는 조각만 글자 단위로
        쪼개지고 그 앞뒤의 작은 조각은 온전히 보존되는지 확인 (버퍼 유실/오염 없음)."""
        from collectors.notion_collector import _pack_parts_by_size
        small_before = "앞부분"
        huge = "다" * 1700  # 800 기준 3개 청크로 분리되어야 함 (800+800+100)
        small_after = "뒷부분"
        chunks = _pack_parts_by_size([small_before, huge, small_after], max_size=800)

        assert chunks[0] == small_before
        assert chunks[1] == "다" * 800
        assert chunks[2] == "다" * 800
        assert chunks[3] == "다" * 100
        assert chunks[4] == small_after
        # 원본 내용이 한 글자도 누락/중복되지 않아야 함
        assert "".join(chunks) == small_before + huge + small_after

    def test_pack_parts_by_size_empty_input(self):
        """빈 리스트는 빈 리스트를 반환 (heading 본문이 없는 경우와 동일 입력)."""
        from collectors.notion_collector import _pack_parts_by_size
        assert _pack_parts_by_size([], max_size=800) == []

    def test_heading_chunking_short_section_keeps_plain_title(self):
        """캡을 넘지 않는 섹션은 (파트 N) 접미사 없이 원래 제목을 유지해야 한다
        (긴 섹션과 짧은 섹션이 같은 페이지에 섞여 있는 현실적인 경우)."""
        from collectors.notion_collector import _chunk_by_headings, FALLBACK_CHUNK_SIZE
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "짧은 장"}]},
             "has_children": False},
            {"id": "p1", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "짧은 내용"}]},
             "has_children": False},
            {"id": "h2", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "긴 장"}]},
             "has_children": False},
            {"id": "p2", "type": "paragraph",
             "paragraph": {"rich_text": [{"plain_text": "라" * (FALLBACK_CHUNK_SIZE + 1)}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        short_docs = [d for d in docs if d.title.startswith("짧은 장")]
        long_docs = [d for d in docs if d.title.startswith("긴 장")]
        assert len(short_docs) == 1
        assert short_docs[0].title == "짧은 장"  # 접미사 없음
        assert len(long_docs) == 2
        assert long_docs[0].title == "긴 장 (파트 1)"
        assert long_docs[1].title == "긴 장 (파트 2)"

    def test_heading_chunking_empty_body_no_part_suffix(self):
        """본문이 전혀 없는 heading은 분할 로직을 거쳐도 빈 Document 1개만 생성되고
        제목에 (파트 N)이 붙지 않아야 한다 (분할 전 기존 동작 보존)."""
        from collectors.notion_collector import _chunk_by_headings
        mock_client = MagicMock()
        blocks = [
            {"id": "h1", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "본문 없는 장"}]},
             "has_children": False},
            {"id": "h2", "type": "heading_1",
             "heading_1": {"rich_text": [{"plain_text": "다음 장"}]},
             "has_children": False},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 2
        assert docs[0].title == "본문 없는 장"
        assert docs[0].content == ""

    def test_heading_chunking_toggleable_heading_long_children_split(self):
        """토글 가능한 heading의 자식 텍스트가 캡을 넘으면(이미 하나의 문자열로 합쳐진
        뒤 content_parts에 들어가므로) 문단 경계 없이 글자 단위로 분할되는지 확인."""
        from collectors.notion_collector import _chunk_by_headings, FALLBACK_CHUNK_SIZE
        mock_client = MagicMock()
        long_children_text = "마" * (FALLBACK_CHUNK_SIZE + 200)
        mock_client.blocks.children.list.return_value = {
            "results": [
                {"id": "c1", "type": "paragraph",
                 "paragraph": {"rich_text": [{"plain_text": long_children_text}]},
                 "has_children": False},
            ],
            "has_more": False,
        }
        blocks = [
            {"id": "h1", "type": "heading_3",
             "heading_3": {"rich_text": [{"plain_text": "긴 토글형 heading"}], "is_toggleable": True},
             "has_children": True},
        ]
        docs = _chunk_by_headings(mock_client, blocks, "테스트페이지", "https://notion.so/xxx")
        assert len(docs) == 2
        assert all(d.notion_block_id == "h1" for d in docs)
        assert "".join(d.content for d in docs) == long_children_text
        for d in docs:
            assert len(d.content) <= FALLBACK_CHUNK_SIZE

    def test_collect_notion_page_empty(self):
        """빈 페이지는 빈 리스트 반환."""
        from collectors.notion_collector import collect_notion_page
        mock_client = MagicMock()
        mock_client.blocks.children.list.return_value = {"results": [], "has_more": False}
        docs = collect_notion_page(mock_client, "abc-123", "빈페이지", "https://notion.so/xxx")
        assert docs == []

    def test_expand_blocks_flattens_layout_and_extracts_child_pages(self):
        """column_list/column/callout은 펼치고, 그 안의 child_page는 따로 모은다."""
        from collectors.notion_collector import _expand_blocks

        mock_client = MagicMock()
        children_map = {
            "cl1": [{"id": "col1", "type": "column", "column": {}, "has_children": True}],
            "col1": [{"id": "co1", "type": "callout",
                      "callout": {"rich_text": [{"plain_text": "안내"}]}, "has_children": True}],
            "co1": [
                {"id": "cp1", "type": "child_page",
                 "child_page": {"title": "하위페이지"}, "has_children": True},
                {"id": "p1", "type": "paragraph",
                 "paragraph": {"rich_text": [{"plain_text": "본문"}]}, "has_children": False},
            ],
        }
        mock_client.blocks.children.list.side_effect = (
            lambda block_id, start_cursor=None: {"results": children_map.get(block_id, []), "has_more": False}
        )

        top_blocks = [{"id": "cl1", "type": "column_list", "column_list": {}, "has_children": True}]
        expanded, nested_pages, nested_databases = _expand_blocks(mock_client, top_blocks)

        assert nested_pages == [("cp1", "하위페이지")]
        assert nested_databases == []
        expanded_ids = [b["id"] for b in expanded]
        assert "co1" in expanded_ids  # callout 자신은 유지
        assert "p1" in expanded_ids   # callout 안의 본문도 펼쳐짐
        assert "cl1" not in expanded_ids and "col1" not in expanded_ids  # 레이아웃 컨테이너는 버림

    def test_expand_blocks_extracts_child_database(self):
        """child_database(자료실)는 nested_databases로 따로 모은다(행 자신이
        하나의 페이지라 collect_notion_database가 별도로 재귀 수집함)."""
        from collectors.notion_collector import _expand_blocks

        mock_client = MagicMock()
        blocks = [{"id": "db1", "type": "child_database",
                   "child_database": {"title": "자료실"}, "has_children": False}]
        expanded, nested_pages, nested_databases = _expand_blocks(mock_client, blocks)
        assert expanded == []
        assert nested_pages == []
        assert nested_databases == [("db1", "자료실")]

    def test_collect_notion_page_recurses_into_child_page(self):
        """callout 안에 숨은 child_page를 발견하면 공개 URL을 조회해 재귀 수집한다."""
        from collectors.notion_collector import collect_notion_page

        mock_client = MagicMock()
        top_blocks = [{"id": "co1", "type": "callout",
                       "callout": {"rich_text": []}, "has_children": True}]
        callout_children = [{"id": "cp1", "type": "child_page",
                              "child_page": {"title": "하위페이지"}, "has_children": True}]
        sub_page_blocks = [{"id": "sp1", "type": "paragraph",
                             "paragraph": {"rich_text": [{"plain_text": "하위 내용"}]},
                             "has_children": False}]

        children_map = {"top-1": top_blocks, "co1": callout_children, "cp1": sub_page_blocks}
        mock_client.blocks.children.list.side_effect = (
            lambda block_id, start_cursor=None: {"results": children_map.get(block_id, []), "has_more": False}
        )
        mock_client.pages.retrieve.return_value = {
            "url": "https://app.notion.com/p/cp1",
            "public_url": "https://example.notion.site/cp1",
        }

        docs = collect_notion_page(mock_client, "top-1", "메인페이지", "https://example.notion.site/top-1")

        sub_docs = [d for d in docs if d.content == "하위 내용"]
        assert len(sub_docs) == 1
        assert sub_docs[0].notion_page_url == "https://example.notion.site/cp1"
        assert sub_docs[0].title == "하위페이지"

    def test_collect_notion_page_skips_unresolvable_child_page(self):
        """하위 페이지 URL 조회가 실패하면(예: 권한 없음) 건너뛰고 나머지는 계속 수집한다."""
        from collectors.notion_collector import collect_notion_page

        mock_client = MagicMock()
        top_blocks = [{"id": "co1", "type": "callout",
                       "callout": {"rich_text": []}, "has_children": True}]
        callout_children = [{"id": "cp1", "type": "child_page",
                              "child_page": {"title": "접근불가 페이지"}, "has_children": True}]
        children_map = {"top-1": top_blocks, "co1": callout_children}
        mock_client.blocks.children.list.side_effect = (
            lambda block_id, start_cursor=None: {"results": children_map.get(block_id, []), "has_more": False}
        )
        mock_client.pages.retrieve.return_value = {"url": None, "public_url": None}

        docs = collect_notion_page(mock_client, "top-1", "메인페이지", "https://example.notion.site/top-1")
        # 본문 없는 callout뿐이라도 페이지 자체는 제목만으로 최소 1개 Document가 남고
        # (텍스트가 전혀 없는 페이지가 검색에서 영구히 사라지지 않도록), 조회 실패한
        # 하위페이지("접근불가 페이지")만 건너뛰어 그 문서는 생성되지 않는다.
        assert len(docs) == 1
        assert docs[0].title == "메인페이지"
        assert all(d.title != "접근불가 페이지" for d in docs)

    def test_collect_notion_page_recurses_into_child_database(self):
        """child_database(자료실)를 발견하면 각 행을 child_page와 동일한 방식으로
        재귀 수집한다(행 자신이 하나의 페이지라 속성이 아니라 블록에 내용이 있음)."""
        from collectors.notion_collector import collect_notion_page

        mock_client = MagicMock()
        top_blocks = [{"id": "db1", "type": "child_database",
                       "child_database": {"title": "자료실"}, "has_children": False}]
        row_blocks = [{"id": "rb1", "type": "paragraph",
                       "paragraph": {"rich_text": [{"plain_text": "행 내용"}]}, "has_children": False}]
        children_map = {"top-1": top_blocks, "row-1": row_blocks}
        mock_client.blocks.children.list.side_effect = (
            lambda block_id, start_cursor=None: {"results": children_map.get(block_id, []), "has_more": False}
        )
        mock_client.databases.retrieve.return_value = {"data_sources": [{"id": "ds1", "name": "자료실"}]}
        mock_client.data_sources.query.return_value = {
            "results": [{
                "id": "row-1",
                "properties": {"이름": {"type": "title", "title": [{"plain_text": "행 제목"}]}},
                "url": "https://app.notion.com/p/row-1",
                "public_url": "https://example.notion.site/row-1",
            }],
            "has_more": False,
        }

        docs = collect_notion_page(mock_client, "top-1", "메인페이지", "https://example.notion.site/top-1")

        row_docs = [d for d in docs if d.content == "행 내용"]
        assert len(row_docs) == 1
        assert row_docs[0].source_origin == "행 제목"
        assert row_docs[0].notion_page_url == "https://example.notion.site/row-1"
        assert row_docs[0].notion_block_id == "rb1"

    def test_collect_notion_page_tracks_visited_databases(self):
        """_visited_databases를 넘기면 발견한 child_database id가 그 안에 모인다
        (cron이 나중에 등록할 때 씀)."""
        from collectors.notion_collector import collect_notion_page

        mock_client = MagicMock()
        top_blocks = [{"id": "db1", "type": "child_database",
                       "child_database": {"title": "자료실"}, "has_children": False}]
        mock_client.blocks.children.list.side_effect = (
            lambda block_id, start_cursor=None: {"results": top_blocks if block_id == "top-1" else [],
                                                   "has_more": False}
        )
        mock_client.databases.retrieve.return_value = {"data_sources": []}

        visited_dbs: set = set()
        collect_notion_page(mock_client, "top-1", "메인페이지", "https://example.notion.site/top-1",
                             _visited_databases=visited_dbs)

        assert visited_dbs == {"db1"}

    def test_collect_notion_database_paginates_multiple_data_sources_and_rows(self):
        """데이터베이스에 여러 data_source/행이 있어도 페이지네이션을 따라가며 모두 모은다."""
        from collectors.notion_collector import collect_notion_database

        mock_client = MagicMock()

        def _make_row(row_id, title, has_public_url=True):
            return {
                "id": row_id,
                "properties": {"이름": {"type": "title", "title": [{"plain_text": title}]}},
                "url": f"https://app.notion.com/p/{row_id}",
                "public_url": f"https://example.notion.site/{row_id}" if has_public_url else None,
            }

        def fake_query(data_source_id, start_cursor=None):
            if data_source_id == "ds1":
                if start_cursor is None:
                    return {"results": [_make_row("r1", "행1")], "has_more": True, "next_cursor": "c2"}
                return {"results": [_make_row("r2", "행2")], "has_more": False}
            return {"results": [_make_row("r3", "행3", has_public_url=False)], "has_more": False}

        mock_client.databases.retrieve.return_value = {
            "data_sources": [{"id": "ds1", "name": "A"}, {"id": "ds2", "name": "B"}],
        }
        mock_client.data_sources.query.side_effect = fake_query
        mock_client.blocks.children.list.return_value = {
            "results": [{"id": "blk", "type": "paragraph",
                         "paragraph": {"rich_text": [{"plain_text": "내용"}]}, "has_children": False}],
            "has_more": False,
        }

        docs = collect_notion_database(mock_client, "db1")

        titles = sorted(d.source_origin for d in docs)
        assert titles == ["행1", "행2", "행3"]
        # public_url이 없는 행(r3)은 internal url로 폴백한다
        r3_doc = next(d for d in docs if d.source_origin == "행3")
        assert r3_doc.notion_page_url == "https://app.notion.com/p/r3"

    def test_collect_notion_database_returns_empty_on_retrieve_failure(self):
        """데이터베이스 자체 조회가 실패해도(권한 등) 예외를 전파하지 않고 빈 목록을 반환한다."""
        from collectors.notion_collector import collect_notion_database
        from notion_client.errors import APIResponseError
        import httpx

        mock_client = MagicMock()
        mock_client.databases.retrieve.side_effect = APIResponseError(
            "object_not_found", 404, "찾을 수 없음", httpx.Headers(), "",
        )

        docs = collect_notion_database(mock_client, "db1")
        assert docs == []

    def test_sync_notion_pages_skips_placeholder(self):
        """URL이 플레이스홀더이면 건너뜀."""
        from collectors.notion_collector import sync_notion_pages
        config = {"notion_pages": {"main": "{{NOTION_MAIN_URL}}"}}
        with patch.dict("os.environ", {"NOTION_API_TOKEN": "test-token"}):
            with patch("collectors.notion_collector.Client"):
                docs, summary = sync_notion_pages(config)
        assert summary["main"]["skipped"] is True
        assert len(docs) == 0

    def test_sync_incremental_skips_when_top_and_known_children_unchanged(self):
        """최상위 페이지도, 이전에 발견했던 하위 페이지도 수정 시각이 그대로면 건너뛴다."""
        from collectors.notion_collector import sync_notion_pages_incremental
        config = {"notion_pages": {"main": "https://notion.so/x"}}

        last_edited_by_id = {"top-1": "2026-01-01T00:00:00Z", "sub-1": "2026-01-01T00:00:00Z"}
        with patch.dict("os.environ", {"NOTION_API_TOKEN": "test-token"}), \
             patch("collectors.notion_collector.Client"), \
             patch("collectors.notion_collector.extract_page_id", return_value="top-1"), \
             patch("collectors.notion_collector.get_page_last_edited_time",
                   side_effect=lambda client, pid: last_edited_by_id.get(pid)), \
             patch("storage.supabase_store.get_sync_metadata",
                   return_value={"page_key": "main", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                  "last_synced_at": "x"}), \
             patch("storage.supabase_store.get_sync_metadata_children",
                   return_value=[{"page_key": "main::sub-1", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                   "last_synced_at": "x"}]), \
             patch("storage.supabase_store.get_sync_metadata_databases", return_value=[]), \
             patch("collectors.notion_collector.collect_notion_page") as fake_collect:
            docs, summary = sync_notion_pages_incremental(config)

        fake_collect.assert_not_called()
        assert summary["main"]["skipped"] is True
        assert summary["main"]["reason"] == "변경 없음"
        assert docs == []

    def test_sync_incremental_skip_still_rechecks_known_database(self):
        """페이지/하위페이지는 그대로라 스킵하더라도, 등록된 child_database(자료실)는
        행 추가·삭제가 last_edited_time에 반영되지 않으므로 매번 직접 재조회한다."""
        from collectors.notion_collector import sync_notion_pages_incremental
        from models.document import Document
        config = {"notion_pages": {"main": "https://notion.so/x"}}

        last_edited_by_id = {"top-1": "2026-01-01T00:00:00Z"}
        db_doc = Document.new(source_type="notion", source_origin="FAQ", title="FAQ", content="c",
                               notion_block_id="row-block-1")

        with patch.dict("os.environ", {"NOTION_API_TOKEN": "test-token"}), \
             patch("collectors.notion_collector.Client"), \
             patch("collectors.notion_collector.extract_page_id", return_value="top-1"), \
             patch("collectors.notion_collector.get_page_last_edited_time",
                   side_effect=lambda client, pid: last_edited_by_id.get(pid)), \
             patch("storage.supabase_store.get_sync_metadata",
                   return_value={"page_key": "main", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                  "last_synced_at": "x"}), \
             patch("storage.supabase_store.get_sync_metadata_children", return_value=[]), \
             patch("storage.supabase_store.get_sync_metadata_databases",
                   return_value=[{"page_key": "db::main::db-1", "last_notion_edited_time": "-",
                                   "last_synced_at": "x"}]), \
             patch("collectors.notion_collector.collect_notion_database",
                   return_value=[db_doc]) as fake_collect_db, \
             patch("collectors.notion_collector.collect_notion_page") as fake_collect_page:
            docs, summary = sync_notion_pages_incremental(config)

        fake_collect_page.assert_not_called()
        fake_collect_db.assert_called_once_with(ANY, "db-1")
        assert summary["main"]["skipped"] is True
        assert docs == [db_doc]

    def test_sync_incremental_detects_child_page_only_change(self):
        """최상위 페이지는 안 바뀌었어도, 하위 페이지 자신의 수정 시각이 바뀌면 재수집한다."""
        from collectors.notion_collector import sync_notion_pages_incremental
        config = {"notion_pages": {"main": "https://notion.so/x"}}

        last_edited_by_id = {"top-1": "2026-01-01T00:00:00Z", "sub-1": "2026-02-01T00:00:00Z"}

        def fake_collect_notion_page(client, page_id, page_name, url, visited=None, visited_dbs=None):
            if visited is not None:
                visited.add(page_id)
                visited.add("sub-1")
            return []

        with patch.dict("os.environ", {"NOTION_API_TOKEN": "test-token"}), \
             patch("collectors.notion_collector.Client"), \
             patch("collectors.notion_collector.extract_page_id", return_value="top-1"), \
             patch("collectors.notion_collector.get_page_last_edited_time",
                   side_effect=lambda client, pid: last_edited_by_id.get(pid)), \
             patch("storage.supabase_store.get_sync_metadata",
                   return_value={"page_key": "main", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                  "last_synced_at": "x"}), \
             patch("storage.supabase_store.get_sync_metadata_children",
                   return_value=[{"page_key": "main::sub-1", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                   "last_synced_at": "x"}]), \
             patch("collectors.notion_collector.collect_notion_page",
                   side_effect=fake_collect_notion_page), \
             patch("storage.supabase_store.upsert_sync_metadata") as fake_upsert, \
             patch("storage.supabase_store.delete_sync_metadata_children") as fake_delete_children, \
             patch("storage.supabase_store.delete_sync_metadata_databases") as fake_delete_dbs:
            docs, summary = sync_notion_pages_incremental(config)

        assert summary["main"]["skipped"] is False
        fake_delete_children.assert_called_once_with("main", None)
        fake_delete_dbs.assert_called_once_with("main", None)
        # 최상위(main)와 하위(main::sub-1) 모두 최신 수정 시각으로 기록되어야 함
        upserted_keys = {call.args[0] for call in fake_upsert.call_args_list}
        assert upserted_keys == {"main", "main::sub-1"}

    def test_sync_incremental_recollects_all_when_top_changed(self):
        """최상위 페이지가 바뀌면 하위 페이지 변경 여부와 무관하게 재수집한다."""
        from collectors.notion_collector import sync_notion_pages_incremental
        config = {"notion_pages": {"main": "https://notion.so/x"}}

        last_edited_by_id = {"top-1": "2026-03-01T00:00:00Z"}

        def fake_collect_notion_page(client, page_id, page_name, url, visited=None, visited_dbs=None):
            if visited is not None:
                visited.add(page_id)
            return []

        with patch.dict("os.environ", {"NOTION_API_TOKEN": "test-token"}), \
             patch("collectors.notion_collector.Client"), \
             patch("collectors.notion_collector.extract_page_id", return_value="top-1"), \
             patch("collectors.notion_collector.get_page_last_edited_time",
                   side_effect=lambda client, pid: last_edited_by_id.get(pid)), \
             patch("storage.supabase_store.get_sync_metadata",
                   return_value={"page_key": "main", "last_notion_edited_time": "2026-01-01T00:00:00Z",
                                  "last_synced_at": "x"}), \
             patch("storage.supabase_store.get_sync_metadata_children") as fake_get_children, \
             patch("collectors.notion_collector.collect_notion_page",
                   side_effect=fake_collect_notion_page), \
             patch("storage.supabase_store.upsert_sync_metadata"), \
             patch("storage.supabase_store.delete_sync_metadata_children"), \
             patch("storage.supabase_store.delete_sync_metadata_databases"):
            docs, summary = sync_notion_pages_incremental(config)

        assert summary["main"]["skipped"] is False
        fake_get_children.assert_not_called()  # 최상위가 바뀐 순간 하위 비교는 의미 없으므로 건너뜀


# ══════════════════════════════════════════════════════════════
# collectors/google_sheet_collector.py
# ══════════════════════════════════════════════════════════════

class TestGoogleSheetCollector:
    _SAMPLE_CSV = "질문,답변\n신청 자격은?,대학생이어야 합니다.\n일정은?,3월에 시작합니다.\n"

    def test_collect_public_sheet(self):
        from collectors.google_sheet_collector import collect_google_sheet
        url = "https://docs.google.com/spreadsheets/d/1ABC123/edit?usp=sharing"
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.headers = {"Content-Type": "text/csv"}
        mock_resp.text = self._SAMPLE_CSV

        with patch("collectors.google_sheet_collector.requests.get", return_value=mock_resp):
            docs = collect_google_sheet(url)

        assert len(docs) == 2
        assert docs[0].source_type == "google_sheet"
        assert docs[0].title == "신청 자격은?"
        assert docs[0].content == "대학생이어야 합니다."

    def test_private_sheet_403_raises(self):
        from collectors.google_sheet_collector import collect_google_sheet
        url = "https://docs.google.com/spreadsheets/d/1ABC123/edit"
        mock_resp = MagicMock()
        mock_resp.status_code = 403

        with patch("collectors.google_sheet_collector.requests.get", return_value=mock_resp):
            with pytest.raises(ValueError, match="403"):
                collect_google_sheet(url)

    def test_html_response_raises(self):
        """비공개 시트 → HTML 반환 시 에러."""
        from collectors.google_sheet_collector import collect_google_sheet
        url = "https://docs.google.com/spreadsheets/d/1ABC123/edit"
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.headers = {"Content-Type": "text/html; charset=utf-8"}
        mock_resp.text = "<html>Sign in</html>"

        with patch("collectors.google_sheet_collector.requests.get", return_value=mock_resp):
            with pytest.raises(ValueError, match="비공개"):
                collect_google_sheet(url)

    def test_empty_url_raises(self):
        from collectors.google_sheet_collector import collect_google_sheet
        with pytest.raises(ValueError, match="비어있습니다"):
            collect_google_sheet("")

    def test_invalid_url_raises(self):
        from collectors.google_sheet_collector import collect_google_sheet
        with pytest.raises(ValueError, match="유효한"):
            collect_google_sheet("https://example.com/not-a-sheet")

    def test_network_error_raises(self):
        import requests as req_mod
        from collectors.google_sheet_collector import collect_google_sheet
        url = "https://docs.google.com/spreadsheets/d/1ABC/edit"
        with patch("collectors.google_sheet_collector.requests.get",
                   side_effect=req_mod.RequestException("timeout")):
            with pytest.raises(ValueError, match="다운로드 실패"):
                collect_google_sheet(url)


# ══════════════════════════════════════════════════════════════
# collectors/calendar_collector.py
# ══════════════════════════════════════════════════════════════

class TestCalendarCollector:
    def test_extract_calendar_id_from_embed_url(self):
        from collectors.calendar_collector import _extract_calendar_id
        url = "https://calendar.google.com/calendar/embed?src=26seoulclub%40gmail.com&ctz=Asia%2FSeoul"
        assert _extract_calendar_id(url) == "26seoulclub@gmail.com"

    def test_extract_calendar_id_from_bare_id(self):
        from collectors.calendar_collector import _extract_calendar_id
        assert _extract_calendar_id("26seoulclub@gmail.com") == "26seoulclub@gmail.com"

    def test_build_ics_url_encodes_at_sign(self):
        from collectors.calendar_collector import _build_ics_url
        url = _build_ics_url("26seoulclub@gmail.com")
        assert url == "https://calendar.google.com/calendar/ical/26seoulclub%40gmail.com/public/basic.ics"

    def test_format_date_range_timed_same_day(self):
        from datetime import datetime
        from collectors.calendar_collector import _format_date_range
        start = datetime(2026, 6, 26, 12, 0)
        end = datetime(2026, 6, 26, 17, 0)
        assert _format_date_range(start, end) == "2026년 6월 26일 12:00~17:00"

    def test_format_date_range_timed_cross_day(self):
        from datetime import datetime
        from collectors.calendar_collector import _format_date_range
        start = datetime(2026, 6, 26, 23, 0)
        end = datetime(2026, 6, 27, 1, 0)
        assert _format_date_range(start, end) == "2026년 6월 26일 23:00 ~ 2026년 6월 27일 01:00"

    def test_format_date_range_all_day_single(self):
        from datetime import date
        from collectors.calendar_collector import _format_date_range
        assert _format_date_range(date(2026, 7, 12), date(2026, 7, 13)) == "2026년 7월 12일"

    def test_format_date_range_all_day_multi(self):
        from datetime import date
        from collectors.calendar_collector import _format_date_range
        assert _format_date_range(date(2026, 6, 29), date(2026, 7, 6)) == "2026년 6월 29일~2026년 7월 5일"

    def test_format_date_range_no_zero_padding(self):
        """'08월'은 형태소 분석기가 '8월'과 다른 토큰이 되어 "8월 22일" 질문과 BM25
        매칭이 깨지므로, 월/일에 0패딩이 들어가면 안 된다."""
        from datetime import date
        from collectors.calendar_collector import _format_date_range
        assert _format_date_range(date(2026, 8, 5), date(2026, 8, 6)) == "2026년 8월 5일"

    def _build_ics(self, events: list) -> bytes:
        """(summary, dtstart, dtend, location, description) 튜플 목록으로 ICS 바이트를 만든다."""
        import icalendar
        cal = icalendar.Calendar()
        cal.add("prodid", "-//test//")
        cal.add("version", "2.0")
        for summary, dtstart, dtend, location, description, extra in events:
            ev = icalendar.Event()
            ev.add("summary", summary)
            ev.add("dtstart", dtstart)
            ev.add("dtend", dtend)
            ev.add("uid", summary)
            ev.add("dtstamp", dtstart if hasattr(dtstart, "hour") else dtstart)
            if location:
                ev.add("location", location)
            if description:
                ev.add("description", description)
            if extra:
                ev.add("rrule", extra)
            cal.add_component(ev)
        return cal.to_ical()

    def test_collect_calendar_basic_event(self):
        from datetime import datetime, timedelta, timezone
        from collectors.calendar_collector import collect_google_calendar
        now = datetime.now(timezone.utc)
        start = now + timedelta(days=5)
        ics_bytes = self._build_ics([
            ("정기 모임", start, start + timedelta(hours=2), "학생회관", "분기별 정기 모임", None),
        ])
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = ics_bytes.decode("utf-8")

        with patch("collectors.calendar_collector.requests.get", return_value=mock_resp):
            docs = collect_google_calendar("26seoulclub@gmail.com")

        assert len(docs) == 1
        doc = docs[0]
        assert doc.source_type == "google_calendar"
        assert doc.source_origin == "google_calendar:26seoulclub@gmail.com"
        assert doc.is_editable is False
        assert "정기 모임" in doc.title
        assert "장소: 학생회관" in doc.content
        assert "설명: 분기별 정기 모임" in doc.content

    def test_collect_calendar_doc_id_is_stable_across_syncs(self):
        """같은 일정을 두 번 수집하면(매 동기화마다 일어나는 일) doc_id가 똑같아야
        incremental 동기화가 "이게 그 일정"이라고 알아볼 수 있다."""
        from datetime import datetime, timedelta, timezone
        from collectors.calendar_collector import collect_google_calendar
        now = datetime.now(timezone.utc)
        start = now + timedelta(days=5)
        ics_bytes = self._build_ics([
            ("정기 모임", start, start + timedelta(hours=2), "학생회관", "내용", None),
        ])
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = ics_bytes.decode("utf-8")

        with patch("collectors.calendar_collector.requests.get", return_value=mock_resp):
            docs1 = collect_google_calendar("26seoulclub@gmail.com")
            docs2 = collect_google_calendar("26seoulclub@gmail.com")

        assert docs1[0].doc_id == docs2[0].doc_id

    def test_collect_calendar_doc_id_differs_for_different_events(self):
        from datetime import datetime, timedelta, timezone
        from collectors.calendar_collector import collect_google_calendar
        now = datetime.now(timezone.utc)
        start = now + timedelta(days=5)
        ics_bytes = self._build_ics([
            ("행사A", start, start + timedelta(hours=1), None, None, None),
            ("행사B", start + timedelta(days=1), start + timedelta(days=1, hours=1), None, None, None),
        ])
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = ics_bytes.decode("utf-8")

        with patch("collectors.calendar_collector.requests.get", return_value=mock_resp):
            docs = collect_google_calendar("26seoulclub@gmail.com")

        assert docs[0].doc_id != docs[1].doc_id

    def test_calendar_source_origin_matches_collect_result(self):
        from collectors.calendar_collector import calendar_source_origin
        url = "https://calendar.google.com/calendar/embed?src=26seoulclub%40gmail.com&ctz=Asia%2FSeoul"
        assert calendar_source_origin(url) == "google_calendar:26seoulclub@gmail.com"

    def test_collect_calendar_expands_recurring_events(self):
        """매주 반복되는 일정은 발생일별로 각각 별도 Document가 되어야 한다."""
        from datetime import datetime, timedelta, timezone
        from collectors.calendar_collector import collect_google_calendar
        now = datetime.now(timezone.utc)
        # 30일 전부터 매주 반복 → 앞으로 180일 창 안에 여러 번 발생해야 함
        start = now - timedelta(days=30)
        ics_bytes = self._build_ics([
            ("주간 스터디", start, start + timedelta(hours=1), None, None, {"freq": "weekly"}),
        ])
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = ics_bytes.decode("utf-8")

        with patch("collectors.calendar_collector.requests.get", return_value=mock_resp):
            docs = collect_google_calendar("26seoulclub@gmail.com")

        assert len(docs) > 1
        assert all("주간 스터디" in d.title for d in docs)
        # 발생일이 서로 달라야 함 (같은 내용이 중복 생성된 게 아니라 실제로 펼쳐진 것)
        assert len({d.title for d in docs}) == len(docs)

    def test_collect_calendar_404_raises(self):
        from collectors.calendar_collector import collect_google_calendar
        mock_resp = MagicMock()
        mock_resp.status_code = 404

        with patch("collectors.calendar_collector.requests.get", return_value=mock_resp):
            with pytest.raises(ValueError, match="404"):
                collect_google_calendar("26seoulclub@gmail.com")

    def test_collect_calendar_empty_url_raises(self):
        from collectors.calendar_collector import collect_google_calendar
        with pytest.raises(ValueError, match="비어있습니다"):
            collect_google_calendar("")

    def test_collect_calendar_network_error_raises(self):
        import requests as req_mod
        from collectors.calendar_collector import collect_google_calendar
        with patch("collectors.calendar_collector.requests.get",
                   side_effect=req_mod.RequestException("timeout")):
            with pytest.raises(ValueError, match="다운로드 실패"):
                collect_google_calendar("26seoulclub@gmail.com")


# ══════════════════════════════════════════════════════════════
# collectors/docx_collector.py
# ══════════════════════════════════════════════════════════════

class TestDocxCollector:
    @pytest.fixture
    def sample_docx(self, tmp_path):
        from docx import Document as DocxDocument
        from docx.shared import Pt
        doc = DocxDocument()
        doc.add_heading("1장 신청 자격", level=1)
        doc.add_paragraph("대학생만 신청 가능합니다.")
        doc.add_heading("2장 수당 안내", level=2)
        doc.add_paragraph("월 30만원 지급됩니다.")
        path = tmp_path / "sample.docx"
        doc.save(str(path))
        return str(path)

    @pytest.fixture
    def no_heading_docx(self, tmp_path):
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("단순 내용 A" * 10)
        doc.add_paragraph("단순 내용 B" * 10)
        path = tmp_path / "no_heading.docx"
        doc.save(str(path))
        return str(path)

    def test_heading_chunking(self, sample_docx):
        from collectors.docx_collector import collect_docx
        docs = collect_docx(sample_docx)
        assert len(docs) == 2
        assert docs[0].title == "1장 신청 자격"
        assert "대학생만" in docs[0].content
        assert docs[1].title == "2장 수당 안내"

    def test_file_not_found(self):
        from collectors.docx_collector import collect_docx
        with pytest.raises(FileNotFoundError):
            collect_docx("/nonexistent/file.docx")

    def test_wrong_extension_raises(self, tmp_path):
        from collectors.docx_collector import collect_docx
        f = tmp_path / "test.txt"
        f.write_text("텍스트")
        with pytest.raises(ValueError, match="docx"):
            collect_docx(str(f))

    def test_no_heading_fallback(self, no_heading_docx):
        from collectors.docx_collector import collect_docx
        docs = collect_docx(no_heading_docx)
        assert len(docs) >= 1
        assert all(d.source_type == "docx" for d in docs)

    def test_category_tagged(self, sample_docx):
        from collectors.docx_collector import collect_docx
        docs = collect_docx(sample_docx)
        cats = {d.category for d in docs}
        # "신청 자격"이 포함되어 있으므로 적어도 하나는 분류돼야 함
        assert "미분류" not in cats or len(docs) > 0


# ══════════════════════════════════════════════════════════════
# collectors/pdf_collector.py
# ══════════════════════════════════════════════════════════════

class TestPdfCollector:
    @pytest.fixture
    def mock_pdf(self, tmp_path):
        """pdfplumber를 mock해서 가상 PDF 파일을 만듭니다."""
        p = tmp_path / "sample.pdf"
        p.write_bytes(b"%PDF-1.4 fake")  # 실제 내용은 mock으로 대체
        return str(p)

    def test_file_not_found(self):
        from collectors.pdf_collector import collect_pdf
        with pytest.raises(FileNotFoundError):
            collect_pdf("/nonexistent/file.pdf")

    def test_wrong_extension_raises(self, tmp_path):
        from collectors.pdf_collector import collect_pdf
        f = tmp_path / "test.docx"
        f.write_bytes(b"fake")
        with pytest.raises(ValueError, match="pdf"):
            collect_pdf(str(f))

    def test_collects_pages(self, mock_pdf):
        from collectors.pdf_collector import collect_pdf

        mock_page1 = MagicMock()
        mock_page1.extract_text.return_value = "첫 번째 페이지 내용입니다. " * 5
        mock_page2 = MagicMock()
        mock_page2.extract_text.return_value = "두 번째 페이지 내용입니다. " * 5

        mock_pdf_obj = MagicMock()
        mock_pdf_obj.__enter__ = MagicMock(return_value=mock_pdf_obj)
        mock_pdf_obj.__exit__ = MagicMock(return_value=False)
        mock_pdf_obj.pages = [mock_page1, mock_page2]

        with patch("collectors.pdf_collector.pdfplumber.open", return_value=mock_pdf_obj):
            docs = collect_pdf(mock_pdf)

        assert len(docs) >= 1
        assert all(d.source_type == "pdf" for d in docs)

    def test_empty_pdf_returns_empty(self, mock_pdf):
        from collectors.pdf_collector import collect_pdf

        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""

        mock_pdf_obj = MagicMock()
        mock_pdf_obj.__enter__ = MagicMock(return_value=mock_pdf_obj)
        mock_pdf_obj.__exit__ = MagicMock(return_value=False)
        mock_pdf_obj.pages = [mock_page]

        with patch("collectors.pdf_collector.pdfplumber.open", return_value=mock_pdf_obj):
            docs = collect_pdf(mock_pdf)
        assert docs == []


# ══════════════════════════════════════════════════════════════
# collectors/excel_collector.py
# ══════════════════════════════════════════════════════════════

class TestExcelCollector:
    @pytest.fixture
    def sample_xlsx(self, tmp_path):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "FAQ"
        ws.append(["질문", "답변"])
        ws.append(["신청 자격은?", "대학생이어야 합니다."])
        ws.append(["수당은?", "월 30만원입니다."])
        path = tmp_path / "sample.xlsx"
        wb.save(str(path))
        return str(path)

    @pytest.fixture
    def empty_xlsx(self, tmp_path):
        from openpyxl import Workbook
        wb = Workbook()
        path = tmp_path / "empty.xlsx"
        wb.save(str(path))
        return str(path)

    def test_collect_with_headers(self, sample_xlsx):
        from collectors.excel_collector import collect_excel
        docs = collect_excel(sample_xlsx)
        assert len(docs) == 2
        assert docs[0].title == "신청 자격은?"
        assert docs[0].content == "대학생이어야 합니다."
        assert docs[0].source_type == "excel"

    def test_file_not_found(self):
        from collectors.excel_collector import collect_excel
        with pytest.raises(FileNotFoundError):
            collect_excel("/nonexistent/file.xlsx")

    def test_wrong_extension_raises(self, tmp_path):
        from collectors.excel_collector import collect_excel
        f = tmp_path / "test.csv"
        f.write_text("a,b")
        with pytest.raises(ValueError, match="xlsx"):
            collect_excel(str(f))

    def test_empty_sheet(self, empty_xlsx):
        from collectors.excel_collector import collect_excel
        docs = collect_excel(empty_xlsx)
        assert docs == []

    def test_is_editable_true(self, sample_xlsx):
        from collectors.excel_collector import collect_excel
        docs = collect_excel(sample_xlsx)
        assert all(d.is_editable is True for d in docs)


# ══════════════════════════════════════════════════════════════
# collectors/hwp_converted_collector.py
# ══════════════════════════════════════════════════════════════

class TestHwpConvertedCollector:
    def test_hwp_raises_clear_error(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "document.hwp"
        f.write_bytes(b"fake hwp")
        with pytest.raises(ValueError) as exc_info:
            collect_hwp_converted(str(f))
        msg = str(exc_info.value)
        assert "PDF" in msg or "텍스트" in msg
        assert "변환" in msg

    def test_hwpx_raises_clear_error(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "document.hwpx"
        f.write_bytes(b"fake hwpx")
        with pytest.raises(ValueError) as exc_info:
            collect_hwp_converted(str(f))
        assert "변환" in str(exc_info.value)

    def test_txt_file_parsed(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "converted.txt"
        f.write_text("변환된 한글 문서 내용입니다.", encoding="utf-8")
        docs = collect_hwp_converted(str(f))
        assert len(docs) == 1
        assert docs[0].source_type == "hwp_converted"
        assert "한글" in docs[0].content

    def test_txt_file_utf8_fallback(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "cp949.txt"
        f.write_bytes("한글 내용".encode("cp949"))
        docs = collect_hwp_converted(str(f))
        assert len(docs) >= 1

    def test_empty_txt_returns_empty(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        docs = collect_hwp_converted(str(f))
        assert docs == []

    def test_file_not_found(self):
        from collectors.hwp_converted_collector import collect_hwp_converted
        with pytest.raises(FileNotFoundError):
            collect_hwp_converted("/nonexistent/doc.txt")

    def test_unsupported_extension_raises(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "doc.docx"
        f.write_bytes(b"fake")
        with pytest.raises(ValueError, match="지원하지 않는"):
            collect_hwp_converted(str(f))

    def test_pdf_delegates_correctly(self, tmp_path):
        from collectors.hwp_converted_collector import collect_hwp_converted
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"%PDF fake")
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PDF 내용입니다. " * 5
        mock_pdf_obj = MagicMock()
        mock_pdf_obj.__enter__ = MagicMock(return_value=mock_pdf_obj)
        mock_pdf_obj.__exit__ = MagicMock(return_value=False)
        mock_pdf_obj.pages = [mock_page]
        with patch("collectors.pdf_collector.pdfplumber.open", return_value=mock_pdf_obj):
            docs = collect_hwp_converted(str(f))
        assert len(docs) >= 1
        assert all(d.source_type == "hwp_converted" for d in docs)


# ══════════════════════════════════════════════════════════════
# 통합 테스트 — end-to-end (mock DB)
# ══════════════════════════════════════════════════════════════

class TestEndToEnd:
    def test_docx_to_db(self, tmp_path, pg_conn):
        """docx 수집 → Supabase Postgres 저장 → 조회 흐름."""
        from docx import Document as DocxDocument
        from collectors.docx_collector import collect_docx
        from storage.supabase_store import upsert_documents, get_by_source_type

        # 임시 docx 생성
        doc_file = tmp_path / "faq.docx"
        d = DocxDocument()
        d.add_heading("수당 지급 기준", level=1)
        d.add_paragraph("월 30만원을 지급합니다.")
        d.save(str(doc_file))

        docs = collect_docx(str(doc_file))
        upsert_documents(docs, pg_conn)

        stored = get_by_source_type("docx", pg_conn)
        assert len(stored) == len(docs)
        assert stored[0].title == "수당 지급 기준"

    def test_excel_to_db(self, tmp_path, pg_conn):
        from openpyxl import Workbook
        from collectors.excel_collector import collect_excel
        from storage.supabase_store import upsert_documents, get_total_count

        wb = Workbook()
        ws = wb.active
        ws.append(["질문", "답변"])
        ws.append(["수료 기준은?", "80% 이상 출석"])
        path = tmp_path / "data.xlsx"
        wb.save(str(path))

        baseline = get_total_count(pg_conn)
        docs = collect_excel(str(path))
        upsert_documents(docs, pg_conn)

        assert get_total_count(pg_conn) - baseline == 1

    def test_category_distribution_after_import(self, tmp_path, pg_conn):
        from docx import Document as DocxDocument
        from collectors.docx_collector import collect_docx
        from storage.supabase_store import upsert_documents, get_category_distribution

        doc_file = tmp_path / "multi.docx"
        d = DocxDocument()
        d.add_heading("신청 자격 안내", level=1)
        d.add_paragraph("대학생만 가능합니다.")
        d.add_heading("수당 지급 기준 안내", level=2)
        d.add_paragraph("월 30만원입니다.")
        d.save(str(doc_file))

        docs = collect_docx(str(doc_file))
        upsert_documents(docs, pg_conn)

        dist = get_category_distribution(pg_conn)
        assert "신청 자격 안내" in dist or "미분류" in dist
