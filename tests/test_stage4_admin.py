"""
4단계(운영 대시보드 + 인증) 테스트.
"""

import io
import json
import sys
import uuid
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from unittest.mock import ANY, MagicMock, patch

import jwt
import pytest
from fastapi.testclient import TestClient

_root = Path(__file__).parent.parent
if str(_root) not in sys.path:
    sys.path.insert(0, str(_root))


_TONE_TEST_VALUES = {k: f"테스트-{k}" for k in
                     ["personality", "language_purity", "vip_consistency", "formality",
                      "channel", "emotional_labor", "persona", "factuality"]}


def _make_log_entry(session_id="s1", failure_cause=None, top_score=0.5, timestamp=None):
    return {
        "log_id": str(uuid.uuid4()), "timestamp": (timestamp or datetime.now(timezone.utc)).isoformat(),
        "session_id": session_id, "question": "테스트 질문", "keywords": ["테스트"],
        "question_category": "미분류", "blocked_by_filter": False, "search_success": failure_cause is None,
        "top_score": top_score, "failure_cause": failure_cause, "situation": None,
        "response_attitude": None, "answer": "테스트 답변", "sentiment_score": 0.0,
        "repeated_count": 0, "matched_doc_ids": [], "deep_link": None,
        "escalated_to_operation_team": failure_cause is not None, "latency_ms": 10,
    }


# ══════════════════════════════════════════════════════════════
# auth.py
# ══════════════════════════════════════════════════════════════

class TestAuth:
    def test_hash_and_verify_password_roundtrip(self):
        from auth import hash_password, verify_password
        h = hash_password("correct-password-123")
        assert verify_password("correct-password-123", h) is True

    def test_verify_password_wrong_password_fails(self):
        from auth import hash_password, verify_password
        h = hash_password("correct-password-123")
        assert verify_password("wrong-password", h) is False

    def test_create_and_decode_access_token_roundtrip(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from auth import create_access_token, decode_access_token
        token = create_access_token("ops@test.com")
        assert decode_access_token(token) == "ops@test.com"

    def test_decode_access_token_expired_raises(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from auth import decode_access_token, JWT_ALGORITHM
        now = datetime.now(timezone.utc)
        expired = jwt.encode(
            {"sub": "ops@test.com", "iat": now - timedelta(hours=2), "exp": now - timedelta(hours=1)},
            "test-secret", algorithm=JWT_ALGORITHM,
        )
        with pytest.raises(jwt.ExpiredSignatureError):
            decode_access_token(expired)

    def test_decode_access_token_wrong_secret_raises(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from auth import JWT_ALGORITHM, decode_access_token
        now = datetime.now(timezone.utc)
        token = jwt.encode(
            {"sub": "x", "iat": now, "exp": now + timedelta(hours=1)},
            "different-secret", algorithm=JWT_ALGORITHM,
        )
        with pytest.raises(jwt.InvalidTokenError):
            decode_access_token(token)

    def test_get_current_operator_missing_header_raises_401(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from fastapi import HTTPException
        from auth import get_current_operator
        with pytest.raises(HTTPException) as exc:
            get_current_operator(None)
        assert exc.value.status_code == 401

    def test_get_current_operator_malformed_header_raises_401(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from fastapi import HTTPException
        from auth import get_current_operator
        with pytest.raises(HTTPException) as exc:
            get_current_operator("NotBearer abc")
        assert exc.value.status_code == 401

    def test_get_current_operator_valid_token_returns_email(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        from auth import create_access_token, get_current_operator
        token = create_access_token("ops@test.com")
        assert get_current_operator("Bearer " + token) == "ops@test.com"


# ══════════════════════════════════════════════════════════════
# storage/admin_store.py
# ══════════════════════════════════════════════════════════════

class TestAdminStore:
    def test_create_and_get_operator(self, pg_conn):
        from storage.admin_store import create_operator, get_operator_by_email
        email = f"op-{uuid.uuid4()}@test.com"
        create_operator(email, "hashed-value", pg_conn)
        row = get_operator_by_email(email, pg_conn)
        assert row["email"] == email
        assert row["password_hash"] == "hashed-value"

    def test_get_operator_by_email_unknown_returns_none(self, pg_conn):
        from storage.admin_store import get_operator_by_email
        assert get_operator_by_email("nobody@test.com", pg_conn) is None

    def test_update_password_changes_hash(self, pg_conn):
        from storage.admin_store import create_operator, update_password, get_operator_by_email
        email = f"op-{uuid.uuid4()}@test.com"
        create_operator(email, "old-hash", pg_conn)
        updated = update_password(email, "new-hash", pg_conn)
        assert updated is True
        assert get_operator_by_email(email, pg_conn)["password_hash"] == "new-hash"

    def test_update_password_unknown_email_returns_false(self, pg_conn):
        from storage.admin_store import update_password
        assert update_password("nobody@test.com", "x", pg_conn) is False


# ══════════════════════════════════════════════════════════════
# storage/settings_store.py
# ══════════════════════════════════════════════════════════════

class TestSettingsStore:
    def _defaults(self):
        return {
            "operation_team": {"name": "팀", "address": "주소", "phone": "000", "email_list": ["a@b.com"], "operating_hours": "9-18"},
            "similarity_threshold": 0.55, "search_weights": {"vector": 0.6, "bm25": 0.4},
            "search_top_k": 5, "repeat_threshold": 2, "min_keywords_for_clarity": 1,
            "max_question_length": 500, "rate_limit_per_minute": 10,
            "tone_elements": _TONE_TEST_VALUES,
        }

    # app_settings는 실제 운영 DB에 항상 정확히 1행만 존재하는 싱글톤 테이블이라,
    # "행이 없는 상태"는 운영 데이터를 지우지 않고는 실제 DB로 재현할 수 없습니다.
    # 그래서 그 분기만 cursor를 mock하여 검증하고, 나머지는 실제 DB에 대해
    # 원래 값을 저장해뒀다가 테스트 후 반드시 복원합니다(운영 설정 오염 방지).

    def test_get_settings_raises_when_no_row(self):
        from storage.settings_store import get_settings
        mock_cursor = MagicMock()
        mock_cursor.fetchone.return_value = None
        mock_conn = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cursor
        with pytest.raises(RuntimeError):
            get_settings(mock_conn)

    def test_seed_default_settings_inserts_when_no_row(self):
        from storage.settings_store import seed_default_settings
        mock_cursor = MagicMock()
        mock_cursor.fetchone.return_value = {"cnt": 0}
        mock_conn = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cursor
        seed_default_settings(self._defaults(), mock_conn)
        insert_call = mock_cursor.execute.call_args_list[-1]
        assert "INSERT INTO app_settings" in insert_call[0][0]

    def test_seed_default_settings_noop_when_row_exists(self, pg_conn):
        from storage.settings_store import seed_default_settings, get_settings
        before = get_settings(pg_conn)
        different = self._defaults()
        different["similarity_threshold"] = 0.01
        seed_default_settings(different, pg_conn)  # 이미 행이 있으므로 무시되어야 함
        assert get_settings(pg_conn)["similarity_threshold"] == before["similarity_threshold"]

    def test_update_settings_partial_simple_key(self, pg_conn):
        from storage.settings_store import get_settings, update_settings
        original = get_settings(pg_conn)
        try:
            updated = update_settings({"similarity_threshold": 0.7}, pg_conn)
            assert updated["similarity_threshold"] == 0.7
            assert updated["search_top_k"] == original["search_top_k"]  # 건드리지 않은 값 유지
        finally:
            update_settings(original, pg_conn)

    def test_update_settings_tone_partial(self, pg_conn):
        from storage.settings_store import get_settings, update_settings
        original = get_settings(pg_conn)
        try:
            updated = update_settings({"tone_elements": {"personality": "새 성격"}}, pg_conn)
            assert updated["tone_elements"]["personality"] == "새 성격"
            assert updated["tone_elements"]["formality"] == original["tone_elements"]["formality"]
        finally:
            update_settings(original, pg_conn)

    def test_update_settings_situation_keywords_partial(self, pg_conn):
        from storage.settings_store import get_settings, update_settings
        original = get_settings(pg_conn)
        try:
            new_keywords = {
                "policy_violation": ["대리 출석", "대신 출석"],
                "escalation_request": ["상담원"],
                "gratitude": ["감사합니다"],
                "simple_rejection": ["안 할래요"],
            }
            updated = update_settings({"situation_keywords": new_keywords}, pg_conn)
            assert updated["situation_keywords"] == new_keywords
            assert updated["forbidden_words"] == original["forbidden_words"]  # 건드리지 않은 값 유지
        finally:
            update_settings(original, pg_conn)

    def test_update_settings_forbidden_words_partial(self, pg_conn):
        from storage.settings_store import get_settings, update_settings
        original = get_settings(pg_conn)
        try:
            new_words = {"profanity": ["욕설1"], "hate_speech": [], "threats": ["협박1"]}
            updated = update_settings({"forbidden_words": new_words}, pg_conn)
            assert updated["forbidden_words"] == new_words
        finally:
            update_settings(original, pg_conn)


# ══════════════════════════════════════════════════════════════
# storage/action_store.py
# ══════════════════════════════════════════════════════════════

class TestActionStore:
    def test_get_status_defaults_to_pending(self, pg_conn):
        from storage.action_store import get_status
        from storage.supabase_store import insert_qa_log
        entry = _make_log_entry()
        insert_qa_log(entry, pg_conn)
        assert get_status(entry["log_id"], pg_conn) == "대기"

    def test_set_status_then_get(self, pg_conn):
        from storage.action_store import set_status, get_status
        from storage.supabase_store import insert_qa_log
        entry = _make_log_entry()
        insert_qa_log(entry, pg_conn)
        set_status(entry["log_id"], "처리중", pg_conn)
        assert get_status(entry["log_id"], pg_conn) == "처리중"

    def test_set_status_invalid_value_raises(self, pg_conn):
        from storage.action_store import set_status
        with pytest.raises(ValueError):
            set_status("nonexistent", "보류", pg_conn)

    def test_get_statuses_batch(self, pg_conn):
        from storage.action_store import set_status, get_statuses
        from storage.supabase_store import insert_qa_log
        e1, e2 = _make_log_entry(), _make_log_entry()
        insert_qa_log(e1, pg_conn)
        insert_qa_log(e2, pg_conn)
        set_status(e1["log_id"], "완료", pg_conn)
        result = get_statuses([e1["log_id"], e2["log_id"]], pg_conn)
        assert result[e1["log_id"]] == "완료"
        assert result[e2["log_id"]] == "대기"

    def test_get_statuses_empty_list(self, pg_conn):
        from storage.action_store import get_statuses
        assert get_statuses([], pg_conn) == {}


# ══════════════════════════════════════════════════════════════
# storage/supabase_store.py — 4단계 추가 집계 함수
# ══════════════════════════════════════════════════════════════

class TestQaLogAggregation:
    def test_get_daily_qa_counts(self, pg_conn):
        from storage.supabase_store import insert_qa_log, get_daily_qa_counts
        insert_qa_log(_make_log_entry(), pg_conn)
        counts = get_daily_qa_counts(30, 0, conn=pg_conn)
        assert sum(c["count"] for c in counts) >= 1

    def test_get_daily_qa_counts_empty(self):
        # 실제 운영 DB의 qa_log에는 이미 데이터가 쌓여 있어 "결과 없음"을
        # 실제 DB로 재현할 수 없으므로 cursor를 mock하여 빈 결과 분기를 검증합니다.
        from storage.supabase_store import get_daily_qa_counts
        mock_cursor = MagicMock()
        mock_cursor.fetchall.return_value = []
        mock_conn = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cursor
        assert get_daily_qa_counts(30, 0, conn=mock_conn) == []

    def test_get_daily_qa_counts_pagination(self, pg_conn):
        from storage.supabase_store import insert_qa_log, get_daily_qa_counts
        insert_qa_log(_make_log_entry(), pg_conn)
        page1 = get_daily_qa_counts(1, 0, conn=pg_conn)
        assert len(page1) == 1

    def test_get_qa_logs_paginated(self, pg_conn):
        from storage.supabase_store import insert_qa_log, get_qa_logs_paginated
        for _ in range(3):
            insert_qa_log(_make_log_entry(), pg_conn)
        page1 = get_qa_logs_paginated(2, 0, conn=pg_conn)
        assert len(page1) == 2

    def test_get_logs_by_failure_causes_filters_correctly(self, pg_conn):
        from storage.supabase_store import insert_qa_log, get_logs_by_failure_causes
        entry_incomplete = _make_log_entry(failure_cause="검색실패")
        entry_unresolved = _make_log_entry(failure_cause="지식DB공백")
        insert_qa_log(entry_incomplete, pg_conn)
        insert_qa_log(entry_unresolved, pg_conn)
        incomplete = get_logs_by_failure_causes(["검색실패", "질문모호성"], limit=200, conn=pg_conn)
        ids = {log["log_id"] for log in incomplete}
        assert entry_incomplete["log_id"] in ids
        assert entry_unresolved["log_id"] not in ids
        assert all(log["failure_cause"] in ("검색실패", "질문모호성") for log in incomplete)

    def test_get_logs_by_failure_causes_excludes_resolved(self, pg_conn):
        # action_status='완료'로 표시된 항목은 목록에서 제외되어야 하지만,
        # qa_log 행 자체는 지워지지 않아 통계(get_failure_cause_counts 등)에는 남습니다.
        from storage.supabase_store import insert_qa_log, get_logs_by_failure_causes
        from storage.action_store import set_status
        entry = _make_log_entry(failure_cause="검색실패")
        insert_qa_log(entry, pg_conn)
        set_status(entry["log_id"], "완료", pg_conn)
        incomplete = get_logs_by_failure_causes(["검색실패", "질문모호성"], limit=200, conn=pg_conn)
        assert entry["log_id"] not in {log["log_id"] for log in incomplete}

    def test_delete_old_qa_logs_removes_only_stale_rows(self, pg_conn):
        from storage.supabase_store import insert_qa_log, delete_old_qa_logs
        old_entry = _make_log_entry(timestamp=datetime.now(timezone.utc) - timedelta(days=400))
        recent_entry = _make_log_entry(timestamp=datetime.now(timezone.utc) - timedelta(days=10))
        insert_qa_log(old_entry, pg_conn)
        insert_qa_log(recent_entry, pg_conn)

        deleted = delete_old_qa_logs(365, pg_conn)

        assert deleted >= 1
        with pg_conn.cursor() as cur:
            cur.execute("SELECT log_id FROM qa_log WHERE log_id = %s", (old_entry["log_id"],))
            assert cur.fetchone() is None
            cur.execute("SELECT log_id FROM qa_log WHERE log_id = %s", (recent_entry["log_id"],))
            assert cur.fetchone() is not None

    def test_delete_old_qa_logs_also_removes_action_status_to_avoid_fk_violation(self, pg_conn):
        from storage.supabase_store import insert_qa_log, delete_old_qa_logs
        from storage.action_store import set_status
        old_entry = _make_log_entry(
            failure_cause="검색실패", timestamp=datetime.now(timezone.utc) - timedelta(days=400),
        )
        insert_qa_log(old_entry, pg_conn)
        set_status(old_entry["log_id"], "완료", pg_conn)

        delete_old_qa_logs(365, pg_conn)  # FK 위반 없이 통과해야 함

        with pg_conn.cursor() as cur:
            cur.execute("SELECT log_id FROM action_status WHERE log_id = %s", (old_entry["log_id"],))
            assert cur.fetchone() is None

    def test_get_failure_cause_counts(self, pg_conn):
        from storage.supabase_store import insert_qa_log, get_failure_cause_counts
        before = get_failure_cause_counts(conn=pg_conn)
        insert_qa_log(_make_log_entry(failure_cause="정책밖요청"), pg_conn)
        counts = get_failure_cause_counts(conn=pg_conn)
        assert counts["정책밖요청"] - before.get("정책밖요청", 0) == 1

    # ── 기간별 조회조건 (날짜 범위 필터, KST 기준) ──────────────────

    def _mock_conn_capturing_sql(self):
        mock_cursor = MagicMock()
        mock_cursor.fetchall.return_value = []
        mock_conn = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cursor
        return mock_conn, mock_cursor

    def test_get_daily_qa_counts_date_filter_builds_kst_where_clause(self):
        from storage.supabase_store import get_daily_qa_counts
        mock_conn, mock_cursor = self._mock_conn_capturing_sql()
        get_daily_qa_counts(30, 0, start_date="2026-06-01", end_date="2026-06-21", conn=mock_conn)
        sql, params = mock_cursor.execute.call_args.args
        assert "Asia/Seoul" in sql
        assert "BETWEEN %s AND %s" in sql
        assert params == ("2026-06-01", "2026-06-21", 30, 0)

    def test_get_daily_qa_counts_no_filter_omits_where_clause(self):
        from storage.supabase_store import get_daily_qa_counts
        mock_conn, mock_cursor = self._mock_conn_capturing_sql()
        get_daily_qa_counts(30, 0, conn=mock_conn)
        sql, params = mock_cursor.execute.call_args.args
        assert "WHERE" not in sql
        assert params == (30, 0)

    def test_get_qa_logs_paginated_date_filter_builds_kst_where_clause(self):
        from storage.supabase_store import get_qa_logs_paginated
        mock_conn, mock_cursor = self._mock_conn_capturing_sql()
        get_qa_logs_paginated(50, 0, start_date="2026-06-01", end_date="2026-06-21", conn=mock_conn)
        sql, params = mock_cursor.execute.call_args.args
        assert "Asia/Seoul" in sql
        assert params == ("2026-06-01", "2026-06-21", 50, 0)

    def test_get_logs_by_failure_causes_date_filter_builds_kst_where_clause(self):
        from storage.supabase_store import get_logs_by_failure_causes
        mock_conn, mock_cursor = self._mock_conn_capturing_sql()
        get_logs_by_failure_causes(
            ["검색실패"], limit=50, offset=0, start_date="2026-06-01", end_date="2026-06-21", conn=mock_conn,
        )
        sql, params = mock_cursor.execute.call_args.args
        assert "Asia/Seoul" in sql
        assert params == (["검색실패"], "2026-06-01", "2026-06-21", 50, 0)

    def test_get_failure_cause_counts_date_filter_builds_kst_where_clause(self):
        from storage.supabase_store import get_failure_cause_counts
        mock_conn, mock_cursor = self._mock_conn_capturing_sql()
        get_failure_cause_counts(start_date="2026-06-01", end_date="2026-06-21", conn=mock_conn)
        sql, params = mock_cursor.execute.call_args.args
        assert "Asia/Seoul" in sql
        assert list(params) == ["2026-06-01", "2026-06-21"]

    def test_get_qa_logs_paginated_date_filter_excludes_out_of_range_real_rows(self, pg_conn):
        """실제 DB로 날짜 범위 필터가 진짜로 행을 걸러내는지 확인합니다. 절대 건수가
        아니라 이번 테스트가 직접 넣은 두 행의 존재/부재로만 판단해 공유 DB의
        기존 데이터와 무관하게 안정적으로 동작합니다."""
        from storage.supabase_store import insert_qa_log, get_qa_logs_paginated
        in_range = _make_log_entry(timestamp=datetime(2031, 6, 15, 3, 0, tzinfo=timezone.utc))
        out_of_range = _make_log_entry(timestamp=datetime(2031, 1, 1, 3, 0, tzinfo=timezone.utc))
        insert_qa_log(in_range, pg_conn)
        insert_qa_log(out_of_range, pg_conn)

        logs = get_qa_logs_paginated(
            limit=200, offset=0, start_date="2031-06-01", end_date="2031-06-30", conn=pg_conn,
        )
        ids = {log["log_id"] for log in logs}
        assert in_range["log_id"] in ids
        assert out_of_range["log_id"] not in ids

    def test_get_daily_qa_counts_buckets_by_kst_day_not_utc_day(self, pg_conn):
        """UTC 15:30(=KST 다음날 00:30)에 들어온 로그는 KST 기준 다음날로 집계돼야
        합니다 — UTC 기준이었다면 전날로 집계됐을 시각입니다."""
        from storage.supabase_store import insert_qa_log, get_daily_qa_counts
        entry = _make_log_entry(timestamp=datetime(2032, 3, 10, 15, 30, tzinfo=timezone.utc))
        insert_qa_log(entry, pg_conn)

        counts = get_daily_qa_counts(
            limit=10, offset=0, start_date="2032-03-11", end_date="2032-03-11", conn=pg_conn,
        )
        assert sum(c["count"] for c in counts) == 1
        assert counts[0]["day"] == "2032-03-11"


# ══════════════════════════════════════════════════════════════
# tone_config.py / tone_matrix.py 외부화
# ══════════════════════════════════════════════════════════════

class TestToneOverride:
    def test_build_brand_tone_guideline_default(self):
        from tone_config import build_brand_tone_guideline, BRAND_TONE_ELEMENTS
        guideline = build_brand_tone_guideline()
        assert BRAND_TONE_ELEMENTS["persona"] in guideline

    def test_build_brand_tone_guideline_override(self):
        from tone_config import build_brand_tone_guideline
        custom = {"personality": "엄격한 말투를 사용합니다."}
        guideline = build_brand_tone_guideline(custom)
        assert "엄격한 말투를 사용합니다." in guideline

    def test_tone_matrix_builder_uses_override(self):
        from tone_matrix import ToneMatrixBuilder, Situation
        builder = ToneMatrixBuilder(tone_elements={"personality": "커스텀 성격입니다."})
        instruction = builder.build_instruction(Situation.NORMAL_RESPONSE)
        assert "커스텀 성격입니다." in instruction

    def test_tone_matrix_builder_default_when_none(self):
        from tone_matrix import ToneMatrixBuilder, Situation
        from tone_config import BRAND_TONE_ELEMENTS
        builder = ToneMatrixBuilder()
        instruction = builder.build_instruction(Situation.NORMAL_RESPONSE)
        assert BRAND_TONE_ELEMENTS["persona"] in instruction

    def test_situation_classifier_uses_keyword_override(self):
        from tone_matrix import SituationClassifier, SituationClassificationInput, Situation
        clf = SituationClassifier(keywords={"policy_violation": ["대신 출석"]})
        inp = SituationClassificationInput(
            question="친구가 대신 출석 체크해줘도 되나요?", keywords=["출석"],
            question_category="미분류", top_result_category="미분류", repeated_count=0,
        )
        assert clf.classify(inp) == Situation.POLICY_VIOLATION

    def test_situation_classifier_override_excludes_file_keywords(self):
        # override를 주면 situation_keywords.json 파일 키워드는 전혀 쓰이지 않아야 함
        from tone_matrix import SituationClassifier, SituationClassificationInput, Situation
        clf = SituationClassifier(keywords={"policy_violation": []})
        inp = SituationClassificationInput(
            question="현금으로 따로 받을 수 있나요?", keywords=["현금"],
            question_category="미분류", top_result_category="미분류", repeated_count=0,
        )
        assert clf.classify(inp) != Situation.POLICY_VIOLATION

    def test_load_forbidden_words_uses_override(self):
        from chatbot_engine import _load_forbidden_words, contains_forbidden_word
        words = _load_forbidden_words({"custom": ["테스트금지어"]})
        assert contains_forbidden_word("테스트금지어 포함된 문장", words) is True
        assert contains_forbidden_word("씨발", words) is False  # override 사용 시 파일 기본값은 무시

    def test_load_forbidden_words_default_when_none(self):
        from chatbot_engine import _load_forbidden_words, contains_forbidden_word
        words = _load_forbidden_words()
        assert contains_forbidden_word("씨발 진짜", words) is True


# ══════════════════════════════════════════════════════════════
# api/admin.py
# ══════════════════════════════════════════════════════════════

class TestAddMonths:
    def test_simple_case(self):
        from api.admin import _add_months
        assert _add_months(date(2026, 1, 15), 3) == date(2026, 4, 15)

    def test_crosses_year_boundary(self):
        from api.admin import _add_months
        assert _add_months(date(2026, 11, 1), 3) == date(2027, 2, 1)

    def test_end_of_month_overflow_clamps_to_last_day(self):
        from api.admin import _add_months
        # 1/31 + 1개월은 2월에 31일이 없으므로 2/28(평년)로 클램프됩니다.
        assert _add_months(date(2026, 1, 31), 1) == date(2026, 2, 28)

    def test_zero_months_returns_same_date(self):
        from api.admin import _add_months
        assert _add_months(date(2026, 6, 21), 0) == date(2026, 6, 21)


class TestValidateDateRange:
    def test_both_none_returns_none_none(self):
        from api.admin import _validate_date_range
        assert _validate_date_range(None, None, 3, "테스트") == (None, None)

    def test_only_start_date_raises_400(self):
        from api.admin import _validate_date_range
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _validate_date_range("2026-06-01", None, 3, "테스트")
        assert exc.value.status_code == 400

    def test_invalid_format_raises_400(self):
        from api.admin import _validate_date_range
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _validate_date_range("2026/06/01", "2026-06-21", 3, "테스트")
        assert exc.value.status_code == 400

    def test_start_after_end_raises_400(self):
        from api.admin import _validate_date_range
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _validate_date_range("2026-06-21", "2026-06-01", 3, "테스트")
        assert exc.value.status_code == 400

    def test_exceeds_max_months_raises_400_with_label(self):
        from api.admin import _validate_date_range
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _validate_date_range("2026-01-01", "2026-12-31", 3, "테스트화면")
        assert exc.value.status_code == 400
        assert "테스트화면" in exc.value.detail
        assert "최대 3개월" in exc.value.detail

    def test_valid_range_within_max_months_passes_through(self):
        from api.admin import _validate_date_range
        result = _validate_date_range("2026-06-01", "2026-06-21", 3, "테스트")
        assert result == ("2026-06-01", "2026-06-21")


class TestAdminApi:
    @pytest.fixture
    def operator(self, pg_conn):
        from auth import hash_password
        from storage.admin_store import create_operator
        email = f"ops-{uuid.uuid4()}@test.com"
        create_operator(email, hash_password("correct-password-123"), pg_conn)
        return email

    @pytest.fixture
    def settings_seeded(self, pg_conn):
        # app_settings는 실제 운영 행이 항상 이미 존재하는 싱글톤 테이블입니다.
        # 테스트가 끝나면 반드시 원래 운영 설정값으로 복원해 운영 데이터를
        # 오염시키지 않습니다(이전에 이 부분을 빠뜨려 실제 app_settings가
        # 테스트 값으로 덮어써진 사고가 있었습니다).
        from storage.settings_store import seed_default_settings, get_settings, update_settings
        test_values = {
            "operation_team": {"name": "팀", "address": "주소", "phone": "000-0000-0000",
                                "email_list": ["a@b.com"], "operating_hours": "평일 9-18시"},
            "similarity_threshold": 0.55, "search_weights": {"vector": 0.6, "bm25": 0.4},
            "search_top_k": 5, "repeat_threshold": 2, "min_keywords_for_clarity": 1,
            "max_question_length": 500, "rate_limit_per_minute": 10,
            "tone_elements": _TONE_TEST_VALUES,
        }
        try:
            original = get_settings(pg_conn)
        except RuntimeError:
            original = None
        if original is None:
            seed_default_settings(test_values, pg_conn)
        else:
            update_settings(test_values, pg_conn)
        yield
        if original is not None:
            update_settings(original, pg_conn)

    @pytest.fixture
    def client(self, monkeypatch):
        monkeypatch.setenv("JWT_SECRET_KEY", "test-secret")
        import api.admin as admin_module
        monkeypatch.setattr(admin_module, "_static_config", {
            "notion_pages": {"faq": "https://notion.so/faq-page"},
        })
        return TestClient(admin_module.app)

    def auth_header(self, email):
        from auth import create_access_token
        return {"Authorization": "Bearer " + create_access_token(email)}

    # ── 로그인/비밀번호 변경 ──────────────────────────────────────

    def test_login_success(self, client, operator):
        res = client.post("/login", json={"email": operator, "password": "correct-password-123"})
        assert res.status_code == 200
        assert res.json()["access_token"]

    def test_login_wrong_password_returns_401(self, client, operator):
        res = client.post("/login", json={"email": operator, "password": "wrong"})
        assert res.status_code == 401

    def test_login_unknown_email_returns_401(self, client):
        res = client.post("/login", json={"email": "nobody@test.com", "password": "x"})
        assert res.status_code == 401

    def test_change_password_requires_auth(self, client):
        res = client.post("/change-password", json={"current_password": "a", "new_password": "bbbbbbbb"})
        assert res.status_code == 401

    def test_change_password_success(self, client, operator):
        res = client.post(
            "/change-password",
            json={"current_password": "correct-password-123", "new_password": "new-password-456"},
            headers=self.auth_header(operator),
        )
        assert res.status_code == 200
        # 새 비밀번호로 다시 로그인되는지 확인
        res2 = client.post("/login", json={"email": operator, "password": "new-password-456"})
        assert res2.status_code == 200

    def test_change_password_wrong_current_password_returns_403(self, client, operator):
        # 401이 아니라 403이어야 합니다: 프론트엔드 api() 헬퍼가 401을 받으면
        # 무조건 로그아웃 처리하는데, 비밀번호 오타는 유효한 세션 안의 입력
        # 오류일 뿐 인증 실패가 아니므로 로그아웃되면 안 됩니다.
        res = client.post(
            "/change-password",
            json={"current_password": "wrong", "new_password": "new-password-456"},
            headers=self.auth_header(operator),
        )
        assert res.status_code == 403

    def test_change_password_too_short_returns_400(self, client, operator):
        res = client.post(
            "/change-password",
            json={"current_password": "correct-password-123", "new_password": "short"},
            headers=self.auth_header(operator),
        )
        assert res.status_code == 400

    # ── 모니터링 ────────────────────────────────────────────────

    def test_daily_counts_requires_auth(self, client):
        assert client.get("/monitoring/daily-counts").status_code == 401

    def test_daily_counts_success(self, client, operator, pg_conn):
        from storage.supabase_store import insert_qa_log
        insert_qa_log(_make_log_entry(), pg_conn)
        res = client.get("/monitoring/daily-counts", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert "daily_counts" in res.json()

    def test_daily_counts_invalid_limit_returns_400(self, client, operator):
        res = client.get("/monitoring/daily-counts?limit=0", headers=self.auth_header(operator))
        assert res.status_code == 400

    def test_qa_logs_invalid_limit_returns_400(self, client, operator):
        res = client.get("/monitoring/qa-logs?limit=0", headers=self.auth_header(operator))
        assert res.status_code == 400

    def test_qa_logs_success_empty(self, client, operator):
        res = client.get("/monitoring/qa-logs", headers=self.auth_header(operator))
        assert res.status_code == 200

    # ── 조치관리 ────────────────────────────────────────────────

    def test_incomplete_answers_filters_correctly(self, client, operator, pg_conn):
        from storage.supabase_store import insert_qa_log
        insert_qa_log(_make_log_entry(failure_cause="검색실패"), pg_conn)
        insert_qa_log(_make_log_entry(failure_cause="지식DB공백"), pg_conn)
        res = client.get("/actions/incomplete", headers=self.auth_header(operator))
        assert res.status_code == 200
        causes = {log["failure_cause"] for log in res.json()["logs"]}
        assert causes <= {"검색실패", "질문모호성"}

    def test_unresolved_answers_filters_correctly(self, client, operator, pg_conn):
        from storage.supabase_store import insert_qa_log
        insert_qa_log(_make_log_entry(failure_cause="정책밖요청"), pg_conn)
        res = client.get("/actions/unresolved", headers=self.auth_header(operator))
        assert res.status_code == 200
        for log in res.json()["logs"]:
            assert log["failure_cause"] in ("지식DB공백", "정책밖요청")

    def test_resolve_action_success(self, client, operator, pg_conn):
        from storage.supabase_store import insert_qa_log
        entry = _make_log_entry(failure_cause="검색실패")
        insert_qa_log(entry, pg_conn)
        res = client.delete(f"/actions/{entry['log_id']}", headers=self.auth_header(operator))
        assert res.status_code == 200

    def test_resolve_action_removes_from_incomplete_list(self, client, operator, pg_conn):
        from storage.supabase_store import insert_qa_log
        entry = _make_log_entry(failure_cause="검색실패")
        insert_qa_log(entry, pg_conn)
        client.delete(f"/actions/{entry['log_id']}", headers=self.auth_header(operator))
        res = client.get("/actions/incomplete", headers=self.auth_header(operator))
        ids = {log["log_id"] for log in res.json()["logs"]}
        assert entry["log_id"] not in ids

    def test_resolve_action_nonexistent_log_id_returns_404(self, client, operator):
        res = client.delete("/actions/nonexistent-log-id", headers=self.auth_header(operator))
        assert res.status_code == 404

    def test_failure_report_passes_through_counts(self, client, operator):
        # 실제 운영 DB에 어떤 원인이 몇 건씩 쌓여있는지는 시점마다 달라지므로
        # (공유 DB), 절대값을 검증하지 않고 엔드포인트가 집계 함수의 결과를
        # 그대로 전달하는지만 확인합니다.
        fake_counts = {"지식DB공백": 1, "검색실패": 2, "질문모호성": 3,
                       "정책밖요청": 4, "API오류": 5}
        with patch("storage.supabase_store.get_failure_cause_counts", return_value=fake_counts):
            res = client.get("/actions/failure-report", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json()["counts"] == fake_counts

    def test_failure_report_defaults_missing_causes_to_zero(self, client, operator):
        with patch("storage.supabase_store.get_failure_cause_counts", return_value={"정책밖요청": 3}):
            res = client.get("/actions/failure-report", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json()["counts"] == {
            "지식DB공백": 0, "검색실패": 0, "질문모호성": 0, "정책밖요청": 3, "API오류": 0,
        }

    # ── 기간별 조회조건 (날짜 범위 필터) ────────────────────────────

    @pytest.mark.parametrize("endpoint,store_fn,extra_args", [
        ("/monitoring/daily-counts", "get_daily_qa_counts", ()),
        ("/monitoring/qa-logs", "get_qa_logs_paginated", ()),
        ("/actions/incomplete", "get_logs_by_failure_causes", (["검색실패", "질문모호성"],)),
        ("/actions/unresolved", "get_logs_by_failure_causes", (["지식DB공백", "정책밖요청"],)),
    ])
    def test_date_range_passed_through_to_store_function(
        self, client, operator, endpoint, store_fn, extra_args,
    ):
        with patch(f"storage.supabase_store.{store_fn}", return_value=[]) as fake_store:
            res = client.get(
                f"{endpoint}?start_date=2026-06-01&end_date=2026-06-21",
                headers=self.auth_header(operator),
            )
        assert res.status_code == 200
        call_args = fake_store.call_args.args
        assert call_args[-2:] == ("2026-06-01", "2026-06-21")

    def test_failure_report_date_range_passed_through(self, client, operator):
        with patch("storage.supabase_store.get_failure_cause_counts", return_value={}) as fake_store:
            res = client.get(
                "/actions/failure-report?start_date=2026-01-01&end_date=2026-06-21",
                headers=self.auth_header(operator),
            )
        assert res.status_code == 200
        fake_store.assert_called_once_with("2026-01-01", "2026-06-21")

    @pytest.mark.parametrize("endpoint,max_months", [
        ("/monitoring/daily-counts", 3),
        ("/monitoring/qa-logs", 1),
        ("/actions/incomplete", 3),
        ("/actions/unresolved", 3),
        ("/actions/failure-report", 12),
    ])
    def test_date_range_exceeding_max_months_returns_400(self, client, operator, endpoint, max_months):
        from api.admin import _add_months
        start = "2026-01-01"
        too_far = _add_months(date(2026, 1, 1), max_months) + timedelta(days=1)
        res = client.get(
            f"{endpoint}?start_date={start}&end_date={too_far.isoformat()}",
            headers=self.auth_header(operator),
        )
        assert res.status_code == 400
        assert f"최대 {max_months}개월" in res.json()["detail"]

    def test_date_range_only_start_date_returns_400(self, client, operator):
        res = client.get(
            "/monitoring/daily-counts?start_date=2026-06-01", headers=self.auth_header(operator),
        )
        assert res.status_code == 400

    def test_date_range_start_after_end_returns_400(self, client, operator):
        res = client.get(
            "/monitoring/daily-counts?start_date=2026-06-21&end_date=2026-06-01",
            headers=self.auth_header(operator),
        )
        assert res.status_code == 400

    def test_date_range_invalid_format_returns_400(self, client, operator):
        res = client.get(
            "/monitoring/daily-counts?start_date=2026/06/01&end_date=2026-06-21",
            headers=self.auth_header(operator),
        )
        assert res.status_code == 400

    def test_date_range_exactly_at_max_months_succeeds(self, client, operator):
        from api.admin import _add_months
        start = date(2026, 1, 1)
        end = _add_months(start, 3)
        with patch("storage.supabase_store.get_daily_qa_counts", return_value=[]):
            res = client.get(
                f"/monitoring/daily-counts?start_date={start.isoformat()}&end_date={end.isoformat()}",
                headers=self.auth_header(operator),
            )
        assert res.status_code == 200

    # ── 운영설정 ────────────────────────────────────────────────

    def test_get_settings_success(self, client, operator, settings_seeded):
        res = client.get("/settings", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json()["similarity_threshold"] == 0.55

    def test_update_operation_team_success(self, client, operator, settings_seeded):
        res = client.put("/settings/operation-team", headers=self.auth_header(operator), json={
            "name": "새팀", "address": "새주소", "phone": "02-0000-0000",
            "email_list": ["new@test.com"], "operating_hours": "24시간",
        })
        assert res.status_code == 200
        assert res.json()["operation_team"]["name"] == "새팀"

    def test_update_tone_success_reflected_immediately(self, client, operator, settings_seeded):
        res = client.put("/settings/tone", headers=self.auth_header(operator), json={
            "personality": "새로운 성격", "language_purity": "x", "vip_consistency": "x",
            "formality": "x", "channel": "x", "emotional_labor": "x", "persona": "x", "factuality": "x",
        })
        assert res.status_code == 200
        assert res.json()["tone_elements"]["personality"] == "새로운 성격"

    def test_update_situation_keywords_success(self, client, operator, settings_seeded):
        res = client.put("/settings/situation-keywords", headers=self.auth_header(operator), json={
            "policy_violation": ["대신 출석", "대신 출석"],  # 중복 줄 정리 확인
            "escalation_request": ["상담원"], "gratitude": [" 감사합니다 "],  # 공백 정리 확인
            "simple_rejection": ["", "안 할래요", "  "],  # 빈 줄 정리 확인
        })
        assert res.status_code == 200
        body = res.json()
        assert body["situation_keywords"]["policy_violation"] == ["대신 출석"]
        assert body["situation_keywords"]["gratitude"] == ["감사합니다"]
        assert body["situation_keywords"]["simple_rejection"] == ["안 할래요"]

    def test_update_situation_keywords_missing_category_returns_422(self, client, operator, settings_seeded):
        res = client.put("/settings/situation-keywords", headers=self.auth_header(operator), json={
            "policy_violation": ["대신 출석"],
        })
        assert res.status_code == 422

    def test_update_forbidden_words_success(self, client, operator, settings_seeded):
        res = client.put("/settings/forbidden-words", headers=self.auth_header(operator), json={
            "profanity": ["욕설1"], "hate_speech": [], "threats": ["협박1"],
        })
        assert res.status_code == 200
        assert res.json()["forbidden_words"]["profanity"] == ["욕설1"]
        assert res.json()["forbidden_words"]["hate_speech"] == []

    def test_update_forbidden_words_reflected_in_engine_immediately(self, client, operator, settings_seeded):
        res = client.put("/settings/forbidden-words", headers=self.auth_header(operator), json={
            "profanity": ["테스트금지어다"], "hate_speech": [], "threats": [],
        })
        assert res.status_code == 200
        from chatbot_engine import _load_forbidden_words, contains_forbidden_word
        from storage.settings_store import get_settings
        words = _load_forbidden_words(get_settings()["forbidden_words"])
        assert contains_forbidden_word("테스트금지어다 포함된 문장", words) is True

    def test_update_api_params_invalid_returns_400(self, client, operator, settings_seeded):
        res = client.put("/settings/api-params", headers=self.auth_header(operator),
                          json={"max_question_length": 0, "rate_limit_per_minute": 10})
        assert res.status_code == 400

    def test_update_api_params_success(self, client, operator, settings_seeded):
        res = client.put("/settings/api-params", headers=self.auth_header(operator),
                          json={"max_question_length": 300, "rate_limit_per_minute": 5})
        assert res.status_code == 200
        assert res.json()["max_question_length"] == 300

    # ── Knowledge Base ──────────────────────────────────────────

    def test_list_documents_success(self, client, operator):
        res = client.get("/kb/documents", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert "documents" in res.json()

    def test_delete_document_notion_source_forbidden(self, client, operator, pg_conn):
        from models.document import Document
        from storage.supabase_store import upsert_document
        doc = Document.new(source_type="notion", source_origin="FAQ", title="T", content="C",
                            notion_block_id="abc", is_editable=False)
        upsert_document(doc, pg_conn)
        res = client.delete(f"/kb/documents/{doc.doc_id}", headers=self.auth_header(operator))
        assert res.status_code == 403

    def test_delete_document_editable_succeeds(self, client, operator, pg_conn):
        from models.document import Document
        from storage.supabase_store import upsert_document
        doc = Document.new(source_type="docx", source_origin="a.docx", title="T", content="C")
        upsert_document(doc, pg_conn)
        res = client.delete(f"/kb/documents/{doc.doc_id}", headers=self.auth_header(operator))
        assert res.status_code == 200

    def test_delete_document_not_found_returns_404(self, client, operator):
        res = client.delete("/kb/documents/nonexistent-id", headers=self.auth_header(operator))
        assert res.status_code == 404

    def test_upload_unsupported_extension_returns_400(self, client, operator):
        res = client.post("/kb/upload", headers=self.auth_header(operator),
                           files={"file": ("test.txt", io.BytesIO(b"hello"), "text/plain")})
        assert res.status_code == 400

    def test_upload_hwp_returns_explicit_message(self, client, operator):
        res = client.post("/kb/upload", headers=self.auth_header(operator),
                           files={"file": ("notes.hwp", io.BytesIO(b"x"), "application/octet-stream")})
        assert res.status_code == 400
        assert "PDF" in res.json()["detail"]

    def test_upload_docx_success(self, client, operator):
        from docx import Document as DocxDocument
        buf = io.BytesIO()
        d = DocxDocument()
        d.add_heading("테스트 제목", level=1)
        d.add_paragraph("테스트 내용입니다.")
        d.save(buf)
        buf.seek(0)
        res = client.post("/kb/upload", headers=self.auth_header(operator),
                           files={"file": ("upload-test.docx", buf,
                                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert res.status_code == 200
        assert res.json()["inserted"] >= 1

    def test_embed_all_no_pending_documents_returns_zero(self, client, operator):
        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            res = client.post("/kb/documents/embed-all", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json() == {"status": "ok", "embedded": 0, "failed": 0}

    def test_embed_all_includes_notion_and_calendar_documents(self, client, operator):
        """Voyage rate limit으로 '지금 갱신' 중 임베딩이 일부 실패해도, 노션/캘린더
        문서를 재수집하지 않고 이 버튼으로 임베딩만 재시도할 수 있어야 한다."""
        from models.document import Document
        notion_doc = Document.new(source_type="notion", source_origin="FAQ", title="T",
                                   content="C", notion_block_id="abc", is_editable=False)
        calendar_doc = Document.new(source_type="google_calendar", source_origin="cal", title="행사",
                                     content="일시: 2026년 7월 1일", is_editable=False)
        fake_provider = MagicMock()
        fake_provider.embed_documents.return_value = [[0.1, 0.2], [0.3, 0.4]]
        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding",
                   return_value=[notion_doc, calendar_doc]), \
             patch("embedding_manager.get_embedding_provider", return_value=fake_provider), \
             patch("storage.supabase_store.update_embeddings_batch") as fake_update:
            res = client.post("/kb/documents/embed-all", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json() == {"status": "ok", "embedded": 2, "failed": 0}
        fake_update.assert_called_once()

    def test_embed_all_embeds_pending_editable_documents(self, client, operator):
        from models.document import Document
        doc = Document.new(source_type="google_sheet", source_origin="시트", title="질문1",
                            content="답변1")
        fake_provider = MagicMock()
        fake_provider.embed_documents.return_value = [[0.1, 0.2]]
        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding",
                   return_value=[doc]), \
             patch("embedding_manager.get_embedding_provider", return_value=fake_provider), \
             patch("storage.supabase_store.update_embeddings_batch") as fake_update:
            res = client.post("/kb/documents/embed-all", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json() == {"status": "ok", "embedded": 1, "failed": 0}
        fake_update.assert_called_once_with([(doc.doc_id, [0.1, 0.2])], ANY, conn=ANY)

    def test_embed_all_batch_failure_counts_as_failed_without_crashing(self, client, operator):
        from models.document import Document
        doc = Document.new(source_type="google_sheet", source_origin="시트", title="질문1",
                            content="답변1")
        fake_provider = MagicMock()
        fake_provider.embed_documents.side_effect = ConnectionError("voyage api down")
        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding",
                   return_value=[doc]), \
             patch("embedding_manager.get_embedding_provider", return_value=fake_provider), \
             patch("storage.supabase_store.update_embeddings_batch") as fake_update:
            res = client.post("/kb/documents/embed-all", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json() == {"status": "ok", "embedded": 0, "failed": 1}
        fake_update.assert_not_called()

    def test_google_sheet_empty_url_returns_400(self, client, operator):
        res = client.post("/kb/google-sheet", headers=self.auth_header(operator), json={"url": ""})
        assert res.status_code == 400

    def test_google_sheet_collector_failure_returns_502(self, client, operator):
        with patch("collectors.google_sheet_collector.collect_google_sheet",
                   side_effect=ConnectionError("network down")):
            res = client.post("/kb/google-sheet", headers=self.auth_header(operator),
                               json={"url": "https://docs.google.com/spreadsheets/d/x"})
        assert res.status_code == 502

    def test_notion_refresh_success(self, client, operator, pg_conn):
        fake_result = {"status": "ok", "total_collected": 1, "inserted": 1,
                        "pages": {"faq": {"page_name": "FAQ", "doc_count": 1, "skipped": False}},
                        "sources": {"FAQ": 1}, "validation": {}}
        with patch("api.sync_notion._perform_sync", return_value=fake_result):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert "FAQ" in res.json()["summary_text"]

    def test_notion_refresh_missing_token_returns_503(self, client, operator):
        with patch("api.sync_notion._perform_sync", side_effect=EnvironmentError("NOTION_API_TOKEN 없음")):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        assert res.status_code == 503

    def test_notion_refresh_generic_failure_returns_502(self, client, operator):
        with patch("api.sync_notion._perform_sync", side_effect=ConnectionError("notion api down")):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        assert res.status_code == 502

    def test_notion_refresh_summary_includes_embedding_count(self, client, operator):
        fake_result = {"status": "ok", "total_collected": 1, "inserted": 1,
                        "pages": {"faq": {"page_name": "FAQ", "doc_count": 1, "skipped": False}},
                        "sources": {"FAQ": 1}, "validation": {}, "embedding": {"embedded": 3, "failed": 1}}
        with patch("api.sync_notion._perform_sync", return_value=fake_result):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        summary = res.json()["summary_text"]
        assert "임베딩 3건 반영" in summary
        assert "임베딩 실패 1건" in summary

    def test_notion_refresh_summary_breaks_down_by_actual_source(self, client, operator):
        """pages는 최상위 페이지 키 기준으로 하위 페이지 건수까지 합산돼 있어서,
        운영자에게는 실제 출처(하위 페이지 포함)별 건수를 따로 보여줘야 한다."""
        fake_result = {
            "status": "ok", "total_collected": 12, "inserted": 12,
            "pages": {"main": {"page_name": "메인페이지", "doc_count": 12, "skipped": False},
                       "faq": {"page_name": "FAQ", "skipped": True, "reason": "URL 미설정"}},
            "sources": {"메인페이지": 1, "활동기관 연계": 1, "참여형 프로그램 안내": 5,
                         "국제정원박람회 리더그룹 및 추가모집": 5},
            "validation": {},
        }
        with patch("api.sync_notion._perform_sync", return_value=fake_result):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        summary = res.json()["summary_text"]
        assert "메인페이지 1건 변경" in summary
        assert "활동기관 연계 1건 변경" in summary
        assert "참여형 프로그램 안내 5건 변경" in summary
        assert "국제정원박람회 리더그룹 및 추가모집 5건 변경" in summary
        # URL 미설정(미사용) 페이지는 매번 똑같이 떠서 의미 없는 노이즈이므로 표시하지 않음
        assert "FAQ" not in summary
        assert "메인페이지 12건 변경" not in summary

    def test_notion_refresh_summary_shows_real_skip_reasons_other_than_unconfigured(self, client, operator):
        """URL 미설정이 아닌 실제 실패(예: 노션 조회 오류)는 계속 보여줘야 한다."""
        fake_result = {
            "status": "ok", "total_collected": 0, "inserted": 0,
            "pages": {"main": {"page_name": "메인페이지", "skipped": True, "reason": "조회 실패: 권한 없음"}},
            "sources": {}, "validation": {},
        }
        with patch("api.sync_notion._perform_sync", return_value=fake_result):
            res = client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        summary = res.json()["summary_text"]
        assert "메인페이지 건너뜀(조회 실패: 권한 없음)" in summary

    def test_notion_last_sync_no_record(self, client, operator):
        # sync_metadata는 한 번이라도 갱신(수동/자동)이 발생하면 영구히 행이 남는
        # 실제 운영 테이블이라 "갱신 기록이 아예 없음"을 실제 DB로 재현할 수
        # 없으므로 get_connection을 mock하여 빈 결과 분기를 검증합니다.
        from unittest.mock import patch as _patch
        mock_cursor = MagicMock()
        mock_cursor.fetchone.return_value = None
        mock_conn = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cursor
        with _patch("storage.supabase_store.get_connection", return_value=mock_conn):
            res = client.get("/kb/notion/last-sync", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert res.json()["last_synced_at"] is None
        mock_conn.close.assert_called_once()

    def test_notion_last_sync_after_manual_refresh(self, client, operator, pg_conn):
        fake_result = {"status": "ok", "total_collected": 0, "inserted": 0, "pages": {}, "validation": {}}
        with patch("api.sync_notion._perform_sync", return_value=fake_result):
            client.post("/kb/notion/refresh", headers=self.auth_header(operator))
        res = client.get("/kb/notion/last-sync", headers=self.auth_header(operator))
        assert res.json()["mode"] == "수동"

    def test_notion_faq_url_configured(self, client, operator):
        res = client.get("/kb/notion-faq-url", headers=self.auth_header(operator))
        assert res.json()["url"] == "https://notion.so/faq-page"

    def test_notion_faq_url_placeholder_returns_none(self, client, operator, monkeypatch):
        import api.admin as admin_module
        monkeypatch.setattr(admin_module, "_static_config", {"notion_pages": {"faq": "{{NOTION_FAQ_URL}}"}})
        res = client.get("/kb/notion-faq-url", headers=self.auth_header(operator))
        assert res.json()["url"] is None

    def test_manual_source_guide_success(self, client, operator):
        res = client.get("/kb/manual-source-guide", headers=self.auth_header(operator))
        assert res.status_code == 200
        assert len(res.json()["guide"]) == 4


class TestPerformSyncEmbeddingBackfill:
    """_perform_sync/_perform_incremental_sync가 노션 동기화 직후 누락된 임베딩을
    자동으로 채우는지 직접(엔드포인트를 거치지 않고) 검증합니다."""

    def _write_config(self, tmp_path):
        config = {"notion_pages": {"main": "https://notion.so/x"}, "embedding_model": "voyage-4"}
        (tmp_path / "config.json").write_text(json.dumps(config), encoding="utf-8")
        return config

    def test_perform_sync_backfills_missing_notion_embeddings(self, tmp_path, monkeypatch):
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        doc = Document.new(source_type="notion", source_origin="main", title="t", content="c")
        fake_provider = MagicMock()
        fake_provider.embed_documents.return_value = [[0.1, 0.2]]
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=([doc], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[]), \
             patch("storage.supabase_store.upsert_documents", return_value=1), \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider", return_value=fake_provider), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[doc]), \
             patch("storage.supabase_store.update_embeddings_batch") as fake_update:
            result = sync_notion_module._perform_sync()

        assert result["embedding"] == {"embedded": 1, "failed": 0}
        fake_update.assert_called_once_with([(doc.doc_id, [0.1, 0.2])], "voyage-4", conn=ANY)

    def test_perform_sync_reports_doc_counts_per_actual_source(self, tmp_path, monkeypatch):
        """pages는 최상위 키 기준 합산이라, 실제 출처(하위 페이지 포함)별 건수는
        별도 sources 필드로 노출해야 summary_text가 정확한 분류를 보여줄 수 있다."""
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        docs = [
            Document.new(source_type="notion", source_origin="메인페이지", title="t", content="c"),
            Document.new(source_type="notion", source_origin="활동기관 연계", title="t", content="c"),
            Document.new(source_type="notion", source_origin="활동기관 연계", title="t2", content="c2"),
        ]
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 3, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=(docs, fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[]), \
             patch("storage.supabase_store.upsert_documents", return_value=3), \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = sync_notion_module._perform_sync()

        assert result["sources"] == {"메인페이지": 1, "활동기관 연계": 2}

    def test_perform_sync_embedding_failure_does_not_break_sync(self, tmp_path, monkeypatch):
        """임베딩 단계가 실패해도(예: VOYAGE_API_KEY 누락) 노션 본문 동기화 자체는 성공해야 함."""
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        doc = Document.new(source_type="notion", source_origin="main", title="t", content="c")
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=([doc], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[]), \
             patch("storage.supabase_store.upsert_documents", return_value=1), \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider",
                   side_effect=EnvironmentError("VOYAGE_API_KEY 없음")):
            result = sync_notion_module._perform_sync()

        assert result["status"] == "ok"
        assert result["inserted"] == 1
        assert result["embedding"] == {"embedded": 0, "failed": 0}

    def test_perform_incremental_sync_backfills_missing_notion_embeddings(self, tmp_path, monkeypatch):
        import api.cron.sync_notion as cron_sync_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(cron_sync_module, "_root", tmp_path)

        doc = Document.new(source_type="notion", source_origin="main", title="t", content="c")
        fake_provider = MagicMock()
        fake_provider.embed_documents.return_value = [[0.3, 0.4]]
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages_incremental",
                   return_value=([doc], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_origins", return_value=[]), \
             patch("storage.supabase_store.upsert_documents", return_value=1), \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("storage.supabase_store.delete_old_qa_logs", return_value=0), \
             patch("embedding_manager.get_embedding_provider", return_value=fake_provider), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[doc]), \
             patch("storage.supabase_store.update_embeddings_batch") as fake_update:
            result = cron_sync_module._perform_incremental_sync()

        assert result["embedding"] == {"embedded": 1, "failed": 0}
        fake_update.assert_called_once_with([(doc.doc_id, [0.3, 0.4])], "voyage-4", conn=ANY)


class TestPerformSyncNotionIncremental:
    """_perform_sync(수동 '지금 갱신')가 내용이 안 바뀐 노션 문서의 임베딩을
    보존하고, 바뀐/신규/삭제만 건드리는지 검증합니다. notion_collector가
    (block_id+part_index) 기준으로 doc_id를 고정으로 만들어주는 것에 의존합니다."""

    def _write_config(self, tmp_path):
        config = {"notion_pages": {"main": "https://notion.so/x"}, "embedding_model": "voyage-4"}
        (tmp_path / "config.json").write_text(json.dumps(config), encoding="utf-8")
        return config

    def test_unchanged_notion_doc_preserves_embedding(self, tmp_path, monkeypatch):
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        unchanged = Document.new(source_type="notion", source_origin="메인페이지",
                                  title="T", content="C", notion_block_id="b1")
        unchanged.doc_id = "fixed-id-1"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=([unchanged], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[unchanged]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = sync_notion_module._perform_sync()

        assert result["inserted"] == 0
        assert result["sources"] == {}
        fake_upsert.assert_not_called()
        fake_clear.assert_not_called()

    def test_changed_notion_doc_upserts_and_clears_embedding(self, tmp_path, monkeypatch):
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        old = Document.new(source_type="notion", source_origin="메인페이지",
                            title="T", content="옛 내용", notion_block_id="b1")
        old.doc_id = "fixed-id-1"
        fresh = Document.new(source_type="notion", source_origin="메인페이지",
                              title="T", content="새 내용", notion_block_id="b1")
        fresh.doc_id = "fixed-id-1"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=([fresh], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[old]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = sync_notion_module._perform_sync()

        assert result["inserted"] == 1
        assert result["sources"] == {"메인페이지": 1}
        fake_upsert.assert_called_once_with([fresh], conn=ANY)
        fake_clear.assert_called_once_with(["fixed-id-1"], conn=ANY)

    def test_removed_notion_page_is_deleted(self, tmp_path, monkeypatch):
        """노션에서 페이지/섹션이 통째로 사라지면(워크스페이스에서 삭제·연결해제),
        더 이상 fresh 목록에 없는 기존 문서는 삭제되어야 한다."""
        import api.sync_notion as sync_notion_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(sync_notion_module, "_root", tmp_path)

        gone = Document.new(source_type="notion", source_origin="삭제된 페이지",
                             title="T", content="C", notion_block_id="b-gone")
        gone.doc_id = "fixed-id-gone"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 0, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages", return_value=([], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_type", return_value=[gone]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.delete_by_doc_ids") as fake_delete, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = sync_notion_module._perform_sync()

        assert result["inserted"] == 0
        fake_upsert.assert_not_called()
        fake_delete.assert_called_once_with(["fixed-id-gone"], conn=ANY)


class TestCronPerformIncrementalSyncNotionIncremental:
    """cron 경로(_perform_incremental_sync)도 수동 '지금 갱신'과 동일하게 chunk
    단위로 변경분만 갱신/임베딩 보존해야 한다. 추가로, cron은 메타데이터상
    변경이 감지된 페이지만 재수집하므로(sync_notion_pages_incremental), 변경되지
    않은 페이지의 기존 문서는 비교 대상에서조차 제외되어야 한다(조회 자체를
    하지 않아 임베딩이 그대로 보존됨)."""

    def _write_config(self, tmp_path):
        config = {"notion_pages": {"main": "https://notion.so/x"}, "embedding_model": "voyage-4"}
        (tmp_path / "config.json").write_text(json.dumps(config), encoding="utf-8")
        return config

    def test_unchanged_notion_doc_preserves_embedding(self, tmp_path, monkeypatch):
        import api.cron.sync_notion as cron_sync_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(cron_sync_module, "_root", tmp_path)

        unchanged = Document.new(source_type="notion", source_origin="메인페이지",
                                  title="T", content="C", notion_block_id="b1")
        unchanged.doc_id = "fixed-id-1"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages_incremental",
                   return_value=([unchanged], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_origins", return_value=[unchanged]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("storage.supabase_store.delete_old_qa_logs", return_value=0), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = cron_sync_module._perform_incremental_sync()

        assert result["inserted"] == 0
        assert result["sources"] == {}
        fake_upsert.assert_not_called()
        fake_clear.assert_not_called()

    def test_changed_notion_doc_upserts_and_clears_embedding(self, tmp_path, monkeypatch):
        import api.cron.sync_notion as cron_sync_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(cron_sync_module, "_root", tmp_path)

        old = Document.new(source_type="notion", source_origin="메인페이지",
                            title="T", content="옛 내용", notion_block_id="b1")
        old.doc_id = "fixed-id-1"
        fresh = Document.new(source_type="notion", source_origin="메인페이지",
                              title="T", content="새 내용", notion_block_id="b1")
        fresh.doc_id = "fixed-id-1"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages_incremental",
                   return_value=([fresh], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_origins", return_value=[old]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("storage.supabase_store.delete_old_qa_logs", return_value=0), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = cron_sync_module._perform_incremental_sync()

        assert result["inserted"] == 1
        assert result["sources"] == {"메인페이지": 1}
        fake_upsert.assert_called_once_with([fresh], conn=ANY)
        fake_clear.assert_called_once_with(["fixed-id-1"], conn=ANY)

    def test_removed_block_within_recollected_page_is_deleted(self, tmp_path, monkeypatch):
        """재수집된 페이지 안에서 사라진 block은 삭제되어야 한다."""
        import api.cron.sync_notion as cron_sync_module
        from models.document import Document

        self._write_config(tmp_path)
        monkeypatch.setattr(cron_sync_module, "_root", tmp_path)

        gone = Document.new(source_type="notion", source_origin="메인페이지",
                             title="T", content="C", notion_block_id="b-gone")
        gone.doc_id = "fixed-id-gone"
        survivor = Document.new(source_type="notion", source_origin="메인페이지",
                                 title="T2", content="C2", notion_block_id="b-keep")
        survivor.doc_id = "fixed-id-keep"
        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 1, "skipped": False}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages_incremental",
                   return_value=([survivor], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_origins", return_value=[gone, survivor]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.delete_by_doc_ids") as fake_delete, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("storage.supabase_store.delete_old_qa_logs", return_value=0), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = cron_sync_module._perform_incremental_sync()

        assert result["inserted"] == 0
        fake_upsert.assert_not_called()
        fake_delete.assert_called_once_with(["fixed-id-gone"], conn=ANY)

    def test_unchanged_page_is_never_queried_for_diff(self, tmp_path, monkeypatch):
        """메타데이터상 변경이 없어 재수집되지 않은 페이지(source_origin이 docs에
        전혀 나타나지 않음)는 get_by_source_origins 호출 범위에도 포함되지 않아야
        한다 — 비교 대상에 넣지 않는 것 자체가 임베딩 보존을 보장하는 방법."""
        import api.cron.sync_notion as cron_sync_module

        self._write_config(tmp_path)
        monkeypatch.setattr(cron_sync_module, "_root", tmp_path)

        fake_summary = {"main": {"page_name": "메인페이지", "doc_count": 0, "skipped": True,
                                  "reason": "변경 없음"}}

        with patch("storage.supabase_store.get_connection", return_value=MagicMock()), \
             patch("collectors.notion_collector.sync_notion_pages_incremental",
                   return_value=([], fake_summary)), \
             patch("storage.supabase_store.initialize_db"), \
             patch("storage.supabase_store.get_by_source_origins", return_value=[]) as fake_get, \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("utils.validators.validate_notion_block_ids", return_value={}), \
             patch("storage.supabase_store.delete_old_qa_logs", return_value=0), \
             patch("embedding_manager.get_embedding_provider", return_value=MagicMock()), \
             patch("storage.supabase_store.get_documents_missing_embedding", return_value=[]):
            result = cron_sync_module._perform_incremental_sync()

        assert result["inserted"] == 0
        fake_get.assert_not_called()
        fake_upsert.assert_not_called()


class TestSyncCalendarsIncremental:
    """_sync_calendars가 내용이 안 바뀐 일정의 임베딩을 보존하고, 바뀐/신규/삭제만
    건드리는지 검증합니다 (캘린더 collector가 doc_id를 고정으로 만들어주는 것에
    의존하므로, 가짜 Document에도 같은 doc_id 규칙을 직접 부여해 둡니다)."""

    def _make_doc(self, doc_id, title, content, source_origin="google_calendar:cal@x.com"):
        from models.document import Document
        d = Document.new(source_type="google_calendar", source_origin=source_origin,
                          title=title, content=content, is_editable=False)
        d.doc_id = doc_id
        return d

    _CAL_URL = "https://calendar.google.com/calendar/embed?src=cal%40x.com"
    _ORIGIN = "google_calendar:cal@x.com"

    def test_unchanged_event_is_left_alone_preserving_embedding(self):
        from api.sync_notion import _sync_calendars
        unchanged = self._make_doc("id-1", "행사A", "내용A", source_origin=self._ORIGIN)

        with patch("collectors.calendar_collector.collect_google_calendar", return_value=[unchanged]), \
             patch("storage.supabase_store.get_by_source_origin", return_value=[unchanged]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("storage.supabase_store.delete_by_doc_ids") as fake_delete:
            docs, sources = _sync_calendars({"google_calendars": [self._CAL_URL]}, conn=MagicMock())

        assert docs == []
        assert sources == {self._ORIGIN: 0}
        fake_upsert.assert_not_called()
        fake_clear.assert_not_called()
        fake_delete.assert_called_once_with([], conn=ANY)

    def test_changed_event_upserts_and_clears_embedding(self):
        from api.sync_notion import _sync_calendars
        old = self._make_doc("id-1", "행사A", "예전 내용", source_origin=self._ORIGIN)
        fresh = self._make_doc("id-1", "행사A", "바뀐 내용", source_origin=self._ORIGIN)

        with patch("collectors.calendar_collector.collect_google_calendar", return_value=[fresh]), \
             patch("storage.supabase_store.get_by_source_origin", return_value=[old]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("storage.supabase_store.delete_by_doc_ids"):
            docs, sources = _sync_calendars({"google_calendars": [self._CAL_URL]}, conn=MagicMock())

        assert docs == [fresh]
        assert sources == {self._ORIGIN: 1}
        fake_upsert.assert_called_once_with([fresh], conn=ANY)
        fake_clear.assert_called_once_with(["id-1"], conn=ANY)

    def test_new_event_upserts_without_clearing_embedding(self):
        """신규 행은 upsert_documents가 만드는 새 행이 이미 embedding=NULL이라
        clear_embeddings를 호출할 필요가 없다."""
        from api.sync_notion import _sync_calendars
        new_doc = self._make_doc("id-new", "신규 행사", "내용", source_origin=self._ORIGIN)

        with patch("collectors.calendar_collector.collect_google_calendar", return_value=[new_doc]), \
             patch("storage.supabase_store.get_by_source_origin", return_value=[]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings") as fake_clear, \
             patch("storage.supabase_store.delete_by_doc_ids"):
            docs, sources = _sync_calendars({"google_calendars": [self._CAL_URL]}, conn=MagicMock())

        assert docs == [new_doc]
        assert sources == {self._ORIGIN: 1}
        fake_upsert.assert_called_once_with([new_doc], conn=ANY)
        fake_clear.assert_not_called()

    def test_removed_event_is_deleted(self):
        """더 이상 캘린더에 없는(취소되었거나 -7일 창 밖으로 밀려난) 일정은, 일정이
        0건으로 수집되는 경우(전부 사라짐)에도 source_origin을 URL에서 바로 계산해
        기존 행을 정리할 수 있어야 한다."""
        from api.sync_notion import _sync_calendars
        gone = self._make_doc("id-gone", "지난 행사", "내용", source_origin=self._ORIGIN)

        with patch("collectors.calendar_collector.collect_google_calendar", return_value=[]), \
             patch("storage.supabase_store.get_by_source_origin", return_value=[gone]), \
             patch("storage.supabase_store.upsert_documents") as fake_upsert, \
             patch("storage.supabase_store.clear_embeddings"), \
             patch("storage.supabase_store.delete_by_doc_ids") as fake_delete:
            docs, sources = _sync_calendars({"google_calendars": [self._CAL_URL]}, conn=MagicMock())

        assert docs == []
        assert sources == {self._ORIGIN: 0}
        fake_upsert.assert_not_called()
        fake_delete.assert_called_once_with(["id-gone"], conn=ANY)

    def test_calendar_collection_failure_is_skipped_not_raised(self):
        from api.sync_notion import _sync_calendars
        with patch("collectors.calendar_collector.collect_google_calendar",
                   side_effect=ValueError("비공개 캘린더")):
            docs, sources = _sync_calendars({"google_calendars": [self._CAL_URL]}, conn=MagicMock())
        assert (docs, sources) == ([], {})


# ══════════════════════════════════════════════════════════════
# create_admin_account.py
# ══════════════════════════════════════════════════════════════

class TestCreateAdminAccount:
    def test_main_creates_account(self, monkeypatch, pg_conn, capsys):
        import create_admin_account
        email = f"cli-{uuid.uuid4()}@test.com"
        monkeypatch.setattr(sys, "argv", ["create_admin_account.py", email, "password1234"])
        create_admin_account.main()
        from storage.admin_store import get_operator_by_email
        assert get_operator_by_email(email, pg_conn) is not None
        assert "생성되었습니다" in capsys.readouterr().out

    def test_main_duplicate_email_exits(self, monkeypatch, pg_conn):
        import create_admin_account
        from storage.admin_store import create_operator
        email = f"cli-dup-{uuid.uuid4()}@test.com"
        create_operator(email, "x", pg_conn)
        monkeypatch.setattr(sys, "argv", ["create_admin_account.py", email, "password1234"])
        with pytest.raises(SystemExit):
            create_admin_account.main()

    def test_main_short_password_exits(self, monkeypatch):
        import create_admin_account
        monkeypatch.setattr(sys, "argv", ["create_admin_account.py", "x@test.com", "short"])
        with pytest.raises(SystemExit):
            create_admin_account.main()

    def test_main_missing_args_exits(self, monkeypatch):
        import create_admin_account
        monkeypatch.setattr(sys, "argv", ["create_admin_account.py"])
        with pytest.raises(SystemExit):
            create_admin_account.main()
